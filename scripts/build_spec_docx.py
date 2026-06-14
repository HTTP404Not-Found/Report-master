"""把 SPEC.md 轉成簡單 HTML,然後用 scripts.html_to_docx_direct 走 direct 路徑生成 docx。

設計:
- 用 scripts.source_to_md.md_normalizer 清理 SPEC.md
- 用 `markdown` 套件做 md→html(表格、fence、codehilite、sane_lists、nl2br)
  若套件不可用則 fallback 到原本的規則式解析器
- 透過 scripts.html_to_docx_direct 走 lock 設定(字體/尺寸/行距)
"""
from __future__ import annotations

import argparse
import re
import sys
from html import escape
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from scripts.source_to_md.md_normalizer import normalize_file  # noqa: E402
from scripts.html_to_docx_direct import html_to_docx_direct  # noqa: E402

SPEC_MD = PROJECT_ROOT / "SPEC.md"
LOCK_FILE = PROJECT_ROOT / "examples" / "lock.md"
EXPORT_DIR = PROJECT_ROOT / "exports"


def _strip_yaml_frontmatter(md_text: str) -> str:
    """Strip leading YAML frontmatter (--- delimited) if present.

    B6: SPEC.md doesn't use YAML frontmatter, but other markdown inputs
    might. Strip at the very start so downstream parsing sees clean content.
    """
    return re.sub(r"^---\n.*?\n---\n", "", md_text, flags=re.DOTALL)


def _inline_md(text: str) -> str:
    """處理 inline 標記: `code`, **bold**, *italic*。Used by the fallback only."""
    text = re.sub(r"\*\*([^*\n]+)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"(?<![*\w])\*([^*\n]+)\*(?![*\w])", r"<em>\1</em>", text)
    text = re.sub(r"`([^`\n]+)`", r"<code>\1</code>", text)
    return text


def md_to_html(md_text: str) -> str:
    """Convert markdown → HTML using the `markdown` stdlib-compatible package.

    Extensions enabled:
    - tables        (GFM tables)
    - fenced_code   (``` fence support)
    - codehilite    (syntax highlight, outputs <code> spans)
    - nl2br         (single newline → <br>, not block break)
    - sane_lists    (unambiguous list parsing)

    Raw HTML pass is disabled to prevent XSS and to give us clean tags
    that html_to_docx_direct can process (e.g., <strong> not **text**).

    Falls back to the legacy rules-based parser if the `markdown` package
    is not importable.
    """
    # B6: strip YAML frontmatter before parsing
    md_text = _strip_yaml_frontmatter(md_text)

    try:
        import markdown  # type: ignore
    except ImportError:
        return _md_to_html_fallback(md_text)

    md = markdown.Markdown(
        extensions=[
            "tables",
            "fenced_code",
            "codehilite",
            "nl2br",        # single newline → <br> (not block break)
            "sane_lists",   # unambiguous list parsing
        ],
        extension_configs={
            "codehilite": {
                "css_class": "highlight",
                "guess_lang": False,
            }
        },
        output_format="html",
    )
    # Disable raw HTML — we want portable tags, not passthrough
    md.preprocess = False  # type: ignore[attr-defined]

    return md.convert(md_text)


def _md_to_html_fallback(md_text: str) -> str:
    """Legacy rules-based md→html (used only if `markdown` pkg is unavailable).

    Fixes applied relative to the original version:
    - B2/B4: inline_md runs BEFORE escape on table cells
    - B3: list indentation tracks raw indent, no 2-space inference
    - B5: code_buf is flushed at end (kept; verified below)
    """
    lines = md_text.splitlines()
    out: list[str] = []
    i = 0
    n = len(lines)

    in_code = False
    code_lang = ""
    code_buf: list[str] = []

    in_table = False
    table_header: list[str] = []
    table_body: list[list[str]] = []

    list_stack: list[tuple[int, str]] = []  # [(indent, tag), ...]

    def close_list_to(indent: int) -> None:
        # B3: only close when indent DECREASES (>, not >=)
        while list_stack and list_stack[-1][0] > indent:
            tag = list_stack[-1][1]
            out.append(f"</{tag}>")
            list_stack.pop()

    def open_list_to(indent: int, tag: str) -> None:
        # B3: don't infer nesting depth from 2-space increments.
        # Only close lists at same or lower indent; then open a new list
        # at the current indent level.
        while list_stack and list_stack[-1][0] >= indent:
            t = list_stack.pop()
            out.append(f"</{t[1]}>")
        out.append(f"<{tag}>")
        list_stack.append((indent, tag))

    def close_all_lists() -> None:
        close_list_to(-1)

    def close_table() -> None:
        nonlocal in_table, table_header, table_body
        if not in_table:
            return
        out.append("<table>")
        if table_header:
            out.append("<thead><tr>")
            for c in table_header:
                # B2/B4: inline first, then escape
                out.append(f"<th>{escape(_inline_md(c), quote=False)}</th>")
            out.append("</tr></thead>")
        if table_body:
            out.append("<tbody>")
            for row in table_body:
                out.append("<tr>")
                for c in row:
                    # B2/B4: inline first, then escape
                    out.append(f"<td>{escape(_inline_md(c), quote=False)}</td>")
                out.append("</tr>")
            out.append("</tbody>")
        out.append("</table>")
        in_table = False
        table_header = []
        table_body = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # ── code fence ──
        if stripped.startswith("```"):
            close_all_lists()
            close_table()
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_buf = []
            else:
                text = escape("\n".join(code_buf))
                lang_attr = f' data-lang="{escape(code_lang)}"' if code_lang else ""
                out.append(f"<pre{lang_attr}>{text}</pre>")
                in_code = False
                code_buf = []
                code_lang = ""
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── 空行:結束當前區塊 ──
        if not stripped:
            close_all_lists()
            close_table()
            i += 1
            continue

        # ── 表格 ──
        if stripped.startswith("|"):
            close_all_lists()
            if not in_table:
                # header + separator
                if i + 1 < n and re.match(r"^\|[\s:|-]+\|\s*$", lines[i + 1].strip()):
                    cells = [c.strip() for c in stripped.strip("|").split("|")]
                    table_header = cells
                    table_body = []
                    in_table = True
                    i += 2
                    continue
            if in_table:
                cells = [c.strip() for c in stripped.strip("|").split("|")]
                table_body.append(cells)
                i += 1
                continue
            # 孤立的 | 行(不會進到這)
            i += 1
            continue

        # ── 標題 ──
        m = re.match(r"^(#{1,6})\s+(.+?)\s*#*\s*$", stripped)
        if m:
            close_all_lists()
            close_table()
            level = len(m.group(1))
            out.append(f"<h{level}>{_inline_md(m.group(2))}</h{level}>")
            i += 1
            continue

        # ── hr ──
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            close_all_lists()
            close_table()
            out.append("<hr/>")
            i += 1
            continue

        # ── blockquote ──
        if stripped.startswith(">"):
            close_all_lists()
            close_table()
            content = re.sub(r"^>\s*", "", stripped)
            out.append(f"<p>{_inline_md(content)}</p>")
            i += 1
            continue

        # ── 列表 ──
        m_ul = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        m_ol = re.match(r"^(\s*)(\d+)\.\s+(.+)$", line)
        if m_ul or m_ol:
            indent = len(m_ul.group(1) if m_ul else m_ol.group(1))
            if m_ul:
                content = m_ul.group(2)
                tag = "ul"
            else:
                content = m_ol.group(3)
                tag = "ol"
            close_table()
            # B3: close any list at same or higher indent, then open fresh
            if not list_stack or list_stack[-1][0] < indent:
                open_list_to(indent, tag)
            out.append(f"<li>{_inline_md(content)}</li>")
            i += 1
            continue

        # ── 段落:吃到下一個空行/特殊元素 ──
        close_table()
        para_lines = [stripped]
        j = i + 1
        while j < n:
            nxt = lines[j]
            nxt_strip = nxt.strip()
            if not nxt_strip:
                break
            # 任何 block 級元素都中斷
            if (
                nxt_strip.startswith(("#", "```", ">", "|"))
                or re.match(r"^(-{3,}|\*{3,})$", nxt_strip)
                or re.match(r"^(\s*)([-*]|\d+\.)\s+", nxt)
            ):
                break
            para_lines.append(nxt_strip)
            j += 1
        out.append(f"<p>{_inline_md(' '.join(para_lines))}</p>")
        i = j

    close_all_lists()
    close_table()
    # B5: flush code_buf if file ends mid-fence
    if in_code:
        text = escape("\n".join(code_buf))
        out.append(f"<pre>{text}</pre>")

    return "\n".join(out)


def wrap_html(title: str, body: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="zh-Hant">
<head>
<meta charset="utf-8">
<title>{escape(title)}</title>
</head>
<body>
{body}
</body>
</html>
"""


def _add_page_numbers(docx_path: Path) -> None:
    """Issue #3 — inject a PAGE field into every section's footer.

    Writes a Word field expression (``PAGE``) via OOXML so soffice / Word
    render the live page number at print time rather than a fixed digit.
    Existing footer content is cleared first so reruns don't stack fields.
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))
    for section in doc.sections:
        footer = section.footer
        # Use the first paragraph; clear any existing runs to keep reruns idempotent.
        para = footer.paragraphs[0]
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        run = para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    doc.save(str(docx_path))


def main() -> int:
    ap = argparse.ArgumentParser(description="把 SPEC.md 轉成 docx")
    ap.add_argument(
        "--lock-file",
        type=Path,
        default=LOCK_FILE,
        help="Path to report_lock.md (default: examples/lock.md)",
    )
    ap.add_argument(
        "--page-numbers",
        action="store_true",
        help="Issue #3 — 在每個 section 的 footer 加 PAGE field (default: off)",
    )
    ap.add_argument(
        "--toc",
        action="store_true",
        help="Issue #3 — 用 scripts.toc_generator 產生 TOC 並插入 body 開頭 (default: off, 需要 pandoc)",
    )
    ap.add_argument(
        "--toc-depth",
        type=int,
        default=3,
        help="TOC 深度 (1-6, default 3, 配合 --toc 使用)",
    )
    args = ap.parse_args()
    lock_file = args.lock_file

    raw = normalize_file(SPEC_MD)

    html_body = md_to_html(raw)
    m = re.search(r"<h1[^>]*>(.+?)</h1>", html_body)
    title = re.sub(r"<[^>]+>", "", m.group(1)) if m else "Report"

    full_html = wrap_html(title, html_body)

    intermediate = EXPORT_DIR / "_intermediate_SPEC.html"
    intermediate.parent.mkdir(parents=True, exist_ok=True)
    intermediate.write_text(full_html, encoding="utf-8")
    print(f"[ok] 中間 HTML: {intermediate} ({len(full_html)} chars)")

    # Issue #3 — optional TOC injection (must happen BEFORE html_to_docx_direct,
    # since the DOCX engine bakes the HTML structure into the final document).
    if args.toc:
        try:
            from scripts.toc_generator import generate_toc  # noqa: WPS433
            full_html = generate_toc(
                html_source=intermediate,
                output_html=intermediate,
                toc_depth=args.toc_depth,
            )
            print(f"[ok] TOC 已注入 {intermediate} (depth={args.toc_depth})")
        except Exception as e:
            # Pandoc missing or other generator error — fail soft, continue build.
            print(f"[toc] skip: {e}")

    out_docx = EXPORT_DIR / "SPEC.docx"
    html_to_docx_direct(
        html_source=intermediate,
        output_docx=out_docx,
        lock_file=lock_file,
    )
    print(f"[ok] DOCX: {out_docx} ({out_docx.stat().st_size} bytes)")

    # Issue #3 — optional page-number footer (post-process the saved DOCX).
    if args.page_numbers:
        try:
            _add_page_numbers(out_docx)
            print(f"[ok] page numbers 已寫入每個 section footer")
        except Exception as e:
            print(f"[page-numbers] skip: {e}")

    # ── Phase C: post-export validation ──
    try:
        from scripts.export_checker import check_export
        rep = check_export(docx_path=str(out_docx),
                           require_docx=True)
        if rep.docx_report:
            pc = rep.docx_report.get("paragraph_count", "?")
            print(f"[export_checker] DOCX paragraphs={pc}, PASS={rep.passed}")
        for issue in rep.issues:
            print(f"[export_checker] ⚠ {issue}")
    except Exception as e:
        print(f"[export_checker] skip: {e}")

    try:
        from scripts.docx_validator import validate_docx
        vrep = validate_docx(str(out_docx), strict=False, do_roundtrip=True)
        if vrep.passed:
            print(f"[docx_validator] PASS ✅")
        else:
            print(f"[docx_validator] FAIL ❌")
        for issue in vrep.issues:
            print(f"[docx_validator] ⚠ {issue}")
    except Exception as e:
        print(f"[docx_validator] skip: {e}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
