# build_template.py — Programmatic Word reference template builder
# 對應 SPEC.md §6.1 R1.1 (DOCX fidelity hardening) + architecture.md §介面定義
# 用途：用 python-docx 從零建立 `report-master-template.docx`，把字體與樣式預載好。
#       產出的 .docx 給 `scripts.html_to_docx` 當 `--reference-doc=` 使用。
#
# 設計：
# - 字體鎖死（CJK=標楷體, Latin=Times New Roman）寫入 styles.xml rFonts 的
#   ascii / hAnsi / eastAsia / cs 四個屬性
# - Normal 樣式: 12pt / 行距 1.5
# - Heading 1/2/3: 22pt / 18pt / 14pt / bold
# - Caption: 10pt / 置中
# - Title: 22pt / bold / 置中（封面用）
# - 封面段落依 --type 選不同 placeholder
#
# CLI：
#   python -m scripts.build_template --output templates/reference/report-master-template.docx
#   python -m scripts.build_template --output /tmp/x.docx --type business
#   python -m scripts.build_template --output /tmp/x.docx --type custom --cover-title "我的封面"

from __future__ import annotations

import argparse
import logging
import sys
from pathlib import Path
from typing import Dict, List, Optional, Union

logger = logging.getLogger(__name__)


# ────────────────────────────────────────────────────────────────────
# Exceptions
# ────────────────────────────────────────────────────────────────────

class BuildTemplateError(Exception):
    """Base for build_template failures."""


# ────────────────────────────────────────────────────────────────────
# Constants — locked defaults (對應 shared-standards.md §4 + docs/report_lock_schema.md)
# ────────────────────────────────────────────────────────────────────

DEFAULT_CJK_FONT = "標楷體"
DEFAULT_LATIN_FONT = "Times New Roman"

# Heading sizes (spec: H1=22, H2=18, H3=14)
HEADING_SIZES_PT = {1: 22, 2: 18, 3: 14}

# Body sizes
BODY_SIZE_PT = 12
CAPTION_SIZE_PT = 10
TITLE_SIZE_PT = 22

# Line spacing
BODY_LINE_SPACING = 1.5

# Cover placeholder texts per --type
COVER_PLACEHOLDERS: Dict[str, List[str]] = {
    "academic": [
        "學術論文範本",
        "",
        "Title: Report-master 學術論文範本",
        "Author: <請填寫作者>",
        "Affiliation: <請填寫研究機構>",
        "Date: YYYY-MM-DD",
        "",
        "（請將此段刪除後插入正式內容）",
    ],
    "business": [
        "商業報告範本",
        "",
        "Report Title: Report-master 商業報告",
        "Department: <請填寫部門>",
        "Report ID: <請填寫報告編號>",
        "Date: YYYY-MM-DD",
        "",
        "（請將此段刪除後插入正式內容）",
    ],
    "spec": [
        "技術規格書範本",
        "",
        "Spec Title: Report-master 技術規格書",
        "Product: <請填寫產品名稱>",
        "Version: v0.1.0",
        "Author: <請填寫作者>",
        "Date: YYYY-MM-DD",
        "",
        "（請將此段刪除後插入正式內容）",
    ],
    "gov": [
        "政府公文範本",
        "",
        "Doc Title: Report-master 政府公文",
        "Agency: <請填寫機關名稱>",
        "Doc No.: <請填寫文號 XXX-XXX-XXX>",
        "Date: YYYY-MM-DD",
        "",
        "（請將此段刪除後插入正式內容）",
    ],
    "custom": [
        "Report-master 自訂範本",
        "",
        "Title: <請填寫封面標題>",
        "Subtitle: <請填寫副標題>",
        "Author: <請填寫作者>",
        "Date: YYYY-MM-DD",
        "",
        "（請將此段刪除後插入正式內容）",
    ],
}


# ────────────────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────────────────

def _has_python_docx() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def _apply_font_to_style(style, cjk_font: str, latin_font: str, size_pt: Optional[int] = None) -> None:
    """Set font properties on a python-docx style.

    Writes rFonts ascii/hAnsi/eastAsia/cs (so Word/LibreOffice/WPS 都吃得到)
    + size (via style.font.size).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # ascii / hAnsi (use python-docx high-level API)
    style.font.name = latin_font
    if size_pt is not None:
        from docx.shared import Pt
        style.font.size = Pt(size_pt)

    # eastAsia + cs (need direct XML)
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        # rPr 通常放在 style 的最後
        style.element.append(rPr)

    # Remove existing rFonts if any
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), latin_font)
    rFonts.set(qn("w:hAnsi"), latin_font)
    rFonts.set(qn("w:eastAsia"), cjk_font)
    rFonts.set(qn("w:cs"), latin_font)
    # hint: w:hint="eastAsia" 讓 Word 對混排更聰明
    rFonts.set(qn("w:hint"), "eastAsia")

    # Insert rFonts at the start of rPr (Word likes rFonts before sz/color)
    rPr.insert(0, rFonts)


def _apply_line_spacing(style, multiplier: float) -> None:
    """Set line spacing (in multiples) on a paragraph style."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = style.element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        # pPr 通常在 rPr 之前
        rPr = style.element.find(qn("w:rPr"))
        if rPr is not None:
            rPr.addprevious(pPr)
        else:
            style.element.append(pPr)

    # Remove existing spacing
    existing_spacing = pPr.find(qn("w:spacing"))
    if existing_spacing is not None:
        pPr.remove(existing_spacing)

    spacing = OxmlElement("w:spacing")
    # line spacing in 240ths-of-a-line; for 1.5x use 360
    spacing.set(qn("w:line"), str(int(240 * multiplier)))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _apply_alignment(style, alignment) -> None:
    """Set paragraph alignment on a paragraph style."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = style.element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        rPr = style.element.find(qn("w:rPr"))
        if rPr is not None:
            rPr.addprevious(pPr)
        else:
            style.element.append(pPr)

    # alignment enum values: WD_ALIGN_PARAGRAPH
    #   LEFT=0 / CENTER=1 / RIGHT=2 / JUSTIFY=3
    align_map = {0: "left", 1: "center", 2: "right", 3: "both"}
    val = align_map.get(int(alignment), "left")

    existing = pPr.find(qn("w:jc"))
    if existing is not None:
        pPr.remove(existing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), val)
    pPr.append(jc)


def _apply_bold(style, bold: bool = True) -> None:
    """Force bold on a style."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style.element.append(rPr)

    existing = rPr.find(qn("w:b"))
    if existing is not None:
        rPr.remove(existing)
    existing_cs = rPr.find(qn("w:bCs"))
    if existing_cs is not None:
        rPr.remove(existing_cs)

    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
        bCs = OxmlElement("w:bCs")
        rPr.append(bCs)


# ────────────────────────────────────────────────────────────────────
# Core build
# ────────────────────────────────────────────────────────────────────

def build(
    output_path: Union[str, Path],
    *,
    type: str = "academic",
    cover_title: Optional[str] = None,
    cover_lines: Optional[List[str]] = None,
    cjk_font: str = DEFAULT_CJK_FONT,
    latin_font: str = DEFAULT_LATIN_FONT,
    body_size_pt: int = BODY_SIZE_PT,
    caption_size_pt: int = CAPTION_SIZE_PT,
    title_size_pt: int = TITLE_SIZE_PT,
    line_spacing: float = BODY_LINE_SPACING,
) -> Path:
    """Build a Word reference DOCX with report-master styles.

    Args:
        output_path: where to write the .docx.
        type: cover placeholder type. One of:
              academic | business | spec | gov | custom.
              Ignored if cover_lines is provided.
        cover_title: optional override for the very first cover paragraph
                     (Title style).
        cover_lines: optional explicit list of cover body lines.
        cjk_font: CJK font name (locked default: 標楷體).
        latin_font: Latin font name (locked default: Times New Roman).
        body_size_pt: Normal style size.
        caption_size_pt: Caption style size.
        title_size_pt: Title style size.
        line_spacing: body line spacing multiplier.

    Returns: Path to the written .docx.

    Raises:
        BuildTemplateError: if python-docx is missing or type is invalid.
    """
    if not _has_python_docx():
        raise BuildTemplateError(
            "python-docx 未安裝。請在 venv 安裝：\n"
            "  .venv/bin/pip install python-docx"
        )

    if type not in COVER_PLACEHOLDERS:
        raise BuildTemplateError(
            f"未知的 type: {type!r}。"
            f"可用：{sorted(COVER_PLACEHOLDERS.keys())}"
        )

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ─── Style: Normal (body) ───
    normal = doc.styles["Normal"]
    _apply_font_to_style(normal, cjk_font=cjk_font, latin_font=latin_font, size_pt=body_size_pt)
    _apply_line_spacing(normal, line_spacing)

    # ─── Style: Title (cover) ───
    title_style = doc.styles["Title"]
    _apply_font_to_style(title_style, cjk_font=cjk_font, latin_font=latin_font, size_pt=title_size_pt)
    _apply_bold(title_style, bold=True)
    _apply_alignment(title_style, WD_ALIGN_PARAGRAPH.CENTER)

    # ─── Style: Heading 1/2/3 ───
    for level in (1, 2, 3):
        style = doc.styles[f"Heading {level}"]
        _apply_font_to_style(style, cjk_font=cjk_font, latin_font=latin_font,
                             size_pt=HEADING_SIZES_PT[level])
        _apply_bold(style, bold=True)

    # ─── Style: Caption ───
    caption_style = doc.styles["Caption"]
    _apply_font_to_style(caption_style, cjk_font=cjk_font, latin_font=latin_font,
                         size_pt=caption_size_pt)
    _apply_alignment(caption_style, WD_ALIGN_PARAGRAPH.CENTER)

    # ─── Cover paragraphs ───
    # First paragraph: Title style
    title_text = cover_title if cover_title is not None else COVER_PLACEHOLDERS[type][0]
    p_title = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p_title.style = doc.styles["Title"]
    p_title.text = ""  # clear any default
    run = p_title.add_run(title_text)
    # Run-level fonts also set (belt-and-suspenders: some viewers ignore style-level eastAsia)
    _apply_run_fonts(run, cjk_font=cjk_font, latin_font=latin_font,
                     size_pt=title_size_pt, bold=True)

    # Body cover lines: use Normal style (centered for visual symmetry on cover)
    lines = cover_lines if cover_lines is not None else COVER_PLACEHOLDERS[type][1:]
    for line in lines:
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        if line:  # skip empty placeholders
            run = p.add_run(line)
            _apply_run_fonts(run, cjk_font=cjk_font, latin_font=latin_font,
                             size_pt=body_size_pt, bold=False)
        # Cover lines are centered
        _apply_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)

    # ─── Add a couple of placeholder Heading / Normal paragraphs so the
    # template is usable as-is. ───
    h1 = doc.add_paragraph(style=doc.styles["Heading 1"])
    h1.add_run("第一章 緒論")

    h2 = doc.add_paragraph(style=doc.styles["Heading 2"])
    h2.add_run("1.1 研究背景")

    body_p = doc.add_paragraph(style=doc.styles["Normal"])
    body_run = body_p.add_run(
        "這是 Normal 樣式的範例段落。同時包含中文與 "
        "Times New Roman Latin text，用以驗證字體 fallback。"
    )
    _apply_run_fonts(body_run, cjk_font=cjk_font, latin_font=latin_font,
                     size_pt=body_size_pt)

    cap = doc.add_paragraph(style=doc.styles["Caption"])
    cap.add_run("Figure 1: 範例圖說（Caption 樣式，10pt 置中）")

    # ─── Save ───
    doc.save(str(out))
    logger.info("Template written: %s (%d bytes)", out, out.stat().st_size)
    return out


def _apply_run_fonts(run, cjk_font: str, latin_font: str,
                     size_pt: int, bold: bool = False) -> None:
    """Force CJK + Latin fonts on a single run (defense in depth)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = latin_font
    run.font.size = Pt(size_pt)
    if bold:
        run.bold = True

    rPr = run._element.get_or_add_rPr()
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), latin_font)
    rFonts.set(qn("w:hAnsi"), latin_font)
    rFonts.set(qn("w:eastAsia"), cjk_font)
    rFonts.set(qn("w:cs"), latin_font)
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.insert(0, rFonts)


def _apply_paragraph_alignment(p, alignment) -> None:
    """Set alignment on a Paragraph object (not style)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = p._p.get_or_add_pPr()
    existing = pPr.find(qn("w:jc"))
    if existing is not None:
        pPr.remove(existing)

    align_map = {0: "left", 1: "center", 2: "right", 3: "both"}
    val = align_map.get(int(alignment), "left")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), val)
    pPr.append(jc)


# ────────────────────────────────────────────────────────────────────
# CLI
# ────────────────────────────────────────────────────────────────────

def _main(argv=None) -> int:
    parser = argparse.ArgumentParser(
        prog="python -m scripts.build_template",
        description=(
            "建立 report-master 參考 DOCX (給 pandoc --reference-doc= 用)。"
            "預設字體：CJK=標楷體 / Latin=Times New Roman。"
        ),
    )
    parser.add_argument(
        "--output", "-o",
        required=True,
        help="輸出 .docx 路徑（會自動建立父目錄）",
    )
    parser.add_argument(
        "--type", "-t",
        default="academic",
        choices=sorted(COVER_PLACEHOLDERS.keys()),
        help="封面 placeholder 類型（default: academic）",
    )
    parser.add_argument(
        "--cover-title",
        default=None,
        help="自訂封面標題（覆寫 --type 的預設封面標題）",
    )
    parser.add_argument(
        "--cover-line",
        action="append",
        default=None,
        help=(
            "自訂封面內文一行（可重複使用；會覆寫 --type 的預設封面內文）。"
            "例如：--cover-line 'Author: wai'"
        ),
    )
    parser.add_argument(
        "--cjk-font",
        default=DEFAULT_CJK_FONT,
        help=f"CJK 字體名稱（default: {DEFAULT_CJK_FONT}）",
    )
    parser.add_argument(
        "--latin-font",
        default=DEFAULT_LATIN_FONT,
        help=f"Latin 字體名稱（default: {DEFAULT_LATIN_FONT}）",
    )
    parser.add_argument(
        "--body-size-pt",
        type=int,
        default=BODY_SIZE_PT,
        help=f"Normal 樣式字級 pt（default: {BODY_SIZE_PT}）",
    )
    parser.add_argument(
        "--line-spacing",
        type=float,
        default=BODY_LINE_SPACING,
        help=f"Normal 樣式行距倍數（default: {BODY_LINE_SPACING}）",
    )

    args = parser.parse_args(argv)

    try:
        out = build(
            output_path=args.output,
            type=args.type,
            cover_title=args.cover_title,
            cover_lines=args.cover_line,
            cjk_font=args.cjk_font,
            latin_font=args.latin_font,
            body_size_pt=args.body_size_pt,
            line_spacing=args.line_spacing,
        )
        size = out.stat().st_size
        print(f"✅ Template written: {out} ({size:,} bytes, type={args.type})")
        return 0
    except BuildTemplateError as e:
        print(f"❌ {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(_main())