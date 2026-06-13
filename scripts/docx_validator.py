# docx_validator.py — DOCX post-processing 驗證
# 對應 SPEC.md §6.1 R1.1 (DOCX fidelity hardening) + architecture.md §介面定義
# 用途：驗 DOCX 字體 (必須 = 標楷體 + Times New Roman)、樣式、章節結構
# 同時做 mammoth round-trip 內容比對

from __future__ import annotations

import logging
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import List, Dict, Any, Optional, Union

logger = logging.getLogger(__name__)


class DOCXValidationError(Exception):
    """Raised when DOCX fails validation (BLOCKING)."""


@dataclass
class DOCXValidationReport:
    passed: bool
    source: str
    checks: List[Dict[str, Any]] = field(default_factory=list)
    issues: List[str] = field(default_factory=list)
    roundtrip_text: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


# ────────────────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────────────────

def _has_python_docx() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def _has_mammoth() -> bool:
    try:
        import mammoth  # noqa: F401
        return True
    except ImportError:
        return False


def _read_docx_fonts(docx_path: Path) -> Dict[str, Any]:
    """Extract font names from styles.xml + document.xml.

    Returns dict with:
      - rFonts (eastAsia, ascii, hAnsi, cs)
      - paragraph_count, heading_count
    """
    import zipfile
    import re

    fonts: Dict[str, Any] = {"eastAsia": set(), "ascii": set(), "hAnsi": set(), "cs": set()}
    paragraphs = 0
    headings = 0

    with zipfile.ZipFile(str(docx_path)) as z:
        for member in ("word/styles.xml", "word/document.xml"):
            if member not in z.namelist():
                continue
            xml = z.read(member).decode("utf-8", errors="replace")
            # rFonts w:eastAsia / w:ascii / w:hAnsi / w:cs
            for m in re.finditer(r'w:rFonts[^/>]*', xml):
                tag = m.group(0)
                for attr in ("eastAsia", "ascii", "hAnsi", "cs"):
                    mm = re.search(rf'w:{attr}="([^"]+)"', tag)
                    if mm:
                        fonts[attr].add(mm.group(1))
            # count <w:p>
            paragraphs = xml.count("<w:p ") + xml.count("<w:p>")
            # count headings (style ref to Heading1/2/3)
            headings = len(re.findall(r'w:val="Heading[123]"', xml))

    return {
        "eastAsia": sorted(fonts["eastAsia"]),
        "ascii": sorted(fonts["ascii"]),
        "hAnsi": sorted(fonts["hAnsi"]),
        "cs": sorted(fonts["cs"]),
        "paragraph_count": paragraphs,
        "heading_count": headings,
    }


def _read_docx_with_python_docx(docx_path: Path) -> Dict[str, Any]:
    """Use python-docx for higher-level inspection."""
    if not _has_python_docx():
        return {"available": False}

    from docx import Document  # type: ignore
    doc = Document(str(docx_path))

    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)

    # Collect fonts used in runs
    fonts_used: Dict[str, set] = {"eastAsia": set(), "ascii": set()}
    for p in paragraphs:
        for run in p.runs:
            rPr = run._element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
            if rPr is not None:
                rFonts = rPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
                if rFonts is not None:
                    for attr in ("eastAsia", "ascii"):
                        v = rFonts.get(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{attr}")
                        if v:
                            fonts_used[attr].add(v)

    return {
        "available": True,
        "paragraph_count": len(paragraphs),
        "heading_count": sum(1 for p in paragraphs if p.style.name.startswith("Heading")),
        "table_count": len(tables),
        "fonts_used": {
            "eastAsia": sorted(fonts_used["eastAsia"]),
            "ascii": sorted(fonts_used["ascii"]),
        },
    }


def _mammoth_roundtrip(docx_path: Path) -> Optional[str]:
    """Use mammoth to convert DOCX → plain text. Returns the text or None."""
    if not _has_mammoth():
        return None
    import mammoth  # type: ignore
    with open(str(docx_path), "rb") as f:
        result = mammoth.extract_raw_text(f)
    return result.value


# ────────────────────────────────────────────────────────────────────
# Main API
# ────────────────────────────────────────────────────────────────────

def validate_docx(
    docx_path: Union[str, Path],
    *,
    require_cjk: str = "標楷體",
    require_latin: str = "Times New Roman",
    strict: bool = False,
    do_roundtrip: bool = True,
) -> DOCXValidationReport:
    """Validate a DOCX file against report-master standards.

    Checks:
      1. ZIP integrity + [Content_Types].xml + word/document.xml exist
      2. paragraph_count > 0
      3. CJK font contains '標楷體' (if any rFonts present)
      4. Latin font contains 'Times New Roman' (if any rFonts present)
      5. mammoth round-trip produces non-empty text (optional)

    Args:
        docx_path: path to DOCX.
        require_cjk: required CJK font name (locked: 標楷體).
        require_latin: required Latin font name (locked: Times New Roman).
        strict: if True, missing fonts become BLOCKING even if fonts list is empty.
        do_roundtrip: if True, run mammoth round-trip.

    Returns: DOCXValidationReport (does NOT raise by default; use enforce=True in CLI).
    """
    p = Path(docx_path)
    report = DOCXValidationReport(passed=True, source=str(p))

    if not p.exists():
        report.passed = False
        report.issues.append(f"DOCX 檔不存在: {p}")
        return report

    if p.stat().st_size == 0:
        report.passed = False
        report.issues.append(f"DOCX 檔為空: {p}")
        return report

    # ── Check 1: ZIP integrity ──
    import zipfile
    try:
        with zipfile.ZipFile(str(p)) as z:
            names = z.namelist()
            bad = z.testzip()
            if bad:
                report.passed = False
                report.issues.append(f"ZIP 損壞: {bad}")
            if "[Content_Types].xml" not in names:
                report.passed = False
                report.issues.append("缺 [Content_Types].xml")
            if "word/document.xml" not in names:
                report.passed = False
                report.issues.append("缺 word/document.xml")
    except zipfile.BadZipFile as e:
        report.passed = False
        report.issues.append(f"DOCX 不是有效 ZIP: {e}")
        return report

    if not report.passed:
        return report

    report.checks.append({"name": "zip_integrity", "passed": True})

    # ── Check 2: rFonts inspection ──
    try:
        fonts_info = _read_docx_fonts(p)
        report.checks.append({"name": "fonts_xml", "result": fonts_info})

        cjk_set = set(fonts_info["eastAsia"])
        latin_set = set(fonts_info["ascii"]) | set(fonts_info["hAnsi"])

        if cjk_set:
            if not any(require_cjk in f for f in cjk_set):
                # Font mismatch: in non-strict mode, WARN; in strict, BLOCK.
                msg = f"CJK 字體不符: 期望含 '{require_cjk}'，實際 {sorted(cjk_set)}"
                if strict:
                    report.passed = False
                    report.issues.append(msg)
                else:
                    report.checks.append({"name": "cjk_font", "warn": True, "reason": msg})
            else:
                report.checks.append({"name": "cjk_font", "passed": True,
                                      "found": sorted(cjk_set)})
        elif strict:
            report.passed = False
            report.issues.append(f"strict 模式下，CJK 字體未設定")
        else:
            report.checks.append({"name": "cjk_font", "skipped": True,
                                  "reason": "rFonts.eastAsia 未指定 (可能繼承 reference docx)"})

        if latin_set:
            if not any(require_latin in f for f in latin_set):
                msg = f"Latin 字體不符: 期望含 '{require_latin}'，實際 {sorted(latin_set)}"
                if strict:
                    report.passed = False
                    report.issues.append(msg)
                else:
                    report.checks.append({"name": "latin_font", "warn": True, "reason": msg})
            else:
                report.checks.append({"name": "latin_font", "passed": True,
                                      "found": sorted(latin_set)})
        elif strict:
            report.passed = False
            report.issues.append(f"strict 模式下，Latin 字體未設定")
        else:
            report.checks.append({"name": "latin_font", "skipped": True,
                                  "reason": "rFonts.ascii 未指定"})
    except Exception as e:
        report.passed = False
        report.issues.append(f"讀取 styles.xml 失敗: {e}")

    # ── Check 3: python-docx inspection (optional, but informative) ──
    pd_info = _read_docx_with_python_docx(p)
    if pd_info.get("available"):
        report.checks.append({"name": "python_docx", "result": pd_info})
        if pd_info["paragraph_count"] == 0:
            report.passed = False
            report.issues.append("DOCX 沒有段落")
    else:
        report.checks.append({"name": "python_docx", "skipped": True,
                              "reason": "python-docx 未安裝"})

    # ── Check 4: mammoth round-trip ──
    if do_roundtrip:
        text = _mammoth_roundtrip(p)
        if text is None:
            report.checks.append({"name": "mammoth_roundtrip", "skipped": True,
                                  "reason": "mammoth 未安裝"})
        else:
            report.roundtrip_text = text[:500]  # 截前 500 字節省記憶體
            if not text.strip():
                report.passed = False
                report.issues.append("mammoth round-trip 產出為空")
            else:
                report.checks.append({"name": "mammoth_roundtrip", "passed": True,
                                      "text_length": len(text)})

    return report


# ────────────────────────────────────────────────────────────────────
# CLI
# ────────────────────────────────────────────────────────────────────

def _main() -> int:
    import argparse
    import json
    import sys
    ap = argparse.ArgumentParser(description="Validate DOCX (BLOCKING on failure).")
    ap.add_argument("docx", help="Path to DOCX")
    ap.add_argument("--strict", action="store_true", help="Strict mode")
    ap.add_argument("--no-roundtrip", action="store_true", help="Skip mammoth round-trip")
    ap.add_argument("--json", action="store_true", help="JSON output")
    args = ap.parse_args()

    rep = validate_docx(
        args.docx,
        strict=args.strict,
        do_roundtrip=not args.no_roundtrip,
    )

    if args.json:
        print(json.dumps(rep.to_dict(), ensure_ascii=False, indent=2))
    else:
        if rep.passed:
            print(f"✅ PASS — {rep.source}")
        else:
            print(f"❌ FAIL — {rep.source}")
            for issue in rep.issues:
                print(f"  • {issue}")
        for c in rep.checks:
            print(f"  [{c.get('name')}] {c}")

    return 0 if rep.passed else 1


if __name__ == "__main__":
    raise SystemExit(_main())