# html_to_docx_direct.py — HTML → DOCX via python-docx (parallel path)
# 對應 SPEC.md §3.1 ADR "DOCX 平行路徑" + tasks.md T2-3
# 用途：繞過 Markdown 中間層，直接以 python-docx 將結構化 HTML 段落生成
#       Word 文件。完全控制字體、段落、表格、圖片，適用於政府公文 / 學術
#       投稿等對格式極敏感的場景。
#
# 主路徑仍是 scripts.html_to_docx.html_to_docx (pandoc)。
# 本路徑透過 report_lock.md 的 `output.docx_engine: python-docx` 啟用。
#
# 流程：
#   1. 讀取 lock 檔（預設 examples/lock.md）→ 取得字體/尺寸/行距/頁面設定
#   2. 用 BeautifulSoup 解析 HTML
#   3. 用 python-docx 建立 Document，依序處理 <h1>..<h6>, <p>, <ul>/<ol>,
#      <table>, <img>, <pre>, <code>, <blockquote>, <hr>
#   4. 套用 lock 設定到 styles.xml
#   5. 產出 .docx
#
# CLI：
#   python -m scripts.html_to_docx_direct --input examples/output_1/report_final.html --output /tmp/test.docx
#   python -m scripts.html_to_docx_direct --input report.html --output out.docx --lock-file my_lock.md
#   python -m scripts.html_to_docx_direct --input report.html --output out.docx --template custom-template.docx

from __future__ import annotations

import argparse
import logging
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

# docx imports (top-level so helpers like _add_table / _add_pre can use Pt)
from docx.shared import Pt

# 允許 `python -m scripts.html_to_docx_direct` 從 project root 跑
_PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

logger = logging.getLogger(__name__)


# ────────────────────────────────────────────────────────────────────
# Exceptions
# ────────────────────────────────────────────────────────────────────

class HTMLToDOCXDirectError(Exception):
    """Base for html_to_docx_direct failures."""


class HTMLToDOCXDirectImportError(HTMLToDOCXDirectError):
    """Required dependency (python-docx / bs4 / lxml) is missing."""


class HTMLToDOCXDirectLockError(HTMLToDOCXDirectError):
    """report_lock.md is missing or invalid."""


# ────────────────────────────────────────────────────────────────────
# Dependency checks
# ────────────────────────────────────────────────────────────────────

def _require_deps() -> None:
    """Verify python-docx, bs4, lxml are available."""
    missing: List[str] = []
    try:
        import docx  # noqa: F401
    except ImportError:
        missing.append("python-docx")
    try:
        from bs4 import BeautifulSoup, NavigableString, Tag  # noqa: F401
    except ImportError:
        missing.append("beautifulsoup4")
    try:
        import lxml  # noqa: F401
    except ImportError:
        missing.append("lxml")
    if missing:
        raise HTMLToDOCXDirectImportError(
            f"html_to_docx_direct 需要以下套件：{', '.join(missing)}。"
            f"請 `pip install python-docx beautifulsoup4 lxml`。"
        )


# ────────────────────────────────────────────────────────────────────
# Lock loading
# ────────────────────────────────────────────────────────────────────

def _load_lock(lock_path: Path) -> Dict[str, Any]:
    """Load report_lock.md via scripts.report_lock. Returns YAML dict.

    Falls back to minimal defaults if lock is missing or invalid
    (graceful degradation: the pipeline must not fail on a missing lock
    when smoke-testing; production callers should validate the lock first).
    """
    fallback: Dict[str, Any] = {
        "fonts": {"cjk": "標楷體", "latin": "Times New Roman"},
        "formatting": {
            "body": {"font_size": 12, "line_spacing": 1.5},
            "h1": {"font_size": 18, "bold": True},
            "h2": {"font_size": 16, "bold": True},
            "h3": {"font_size": 14, "bold": True},
            "caption": {"font_size": 10, "align": "center"},
        },
        "page_size": "A4",
        "margins": {"top": "2.5cm", "bottom": "2.5cm",
                   "left": "3cm", "right": "2cm"},
        "line_spacing": 1.5,
    }
    if not lock_path.exists():
        logger.warning("lock file not found: %s — using defaults", lock_path)
        return fallback
    try:
        from scripts.report_lock import read_and_validate
        return read_and_validate(lock_path)
    except Exception as e:
        logger.warning("lock read failed (%s) — using defaults", e)
        return fallback


# ────────────────────────────────────────────────────────────────────
# Unit parsing (cm, mm, pt → EMU)
# ────────────────────────────────────────────────────────────────────

_CM_RE = re.compile(r"^\s*([0-9.]+)\s*cm\s*$", re.IGNORECASE)
_MM_RE = re.compile(r"^\s*([0-9.]+)\s*mm\s*$", re.IGNORECASE)
_IN_RE = re.compile(r"^\s*([0-9.]+)\s*in(?:ch)?s?\s*$", re.IGNORECASE)
_PT_RE = re.compile(r"^\s*([0-9.]+)\s*pt\s*$", re.IGNORECASE)


def _to_cm(value: Union[str, float, int, None], default: float = 2.5) -> float:
    """Parse a length string ('2.5cm', '25mm', '1in', '12pt') → cm (float)."""
    if value is None:
        return default
    if isinstance(value, (int, float)):
        return float(value)
    s = str(value).strip()
    if not s:
        return default
    m = _CM_RE.match(s)
    if m:
        return float(m.group(1))
    m = _MM_RE.match(s)
    if m:
        return float(m.group(1)) / 10.0
    m = _IN_RE.match(s)
    if m:
        return float(m.group(1)) * 2.54
    m = _PT_RE.match(s)
    if m:
        return float(m.group(1)) * 0.0352778
    # Try as bare number (assume cm)
    try:
        return float(s)
    except ValueError:
        return default


# Page size presets (in cm)
_PAGE_SIZES_CM: Dict[str, tuple] = {
    "A4": (21.0, 29.7),
    "A5": (14.8, 21.0),
    "Letter": (21.59, 27.94),
    "Legal": (21.59, 35.56),
    "JIS-B5": (18.2, 25.7),
}


# ────────────────────────────────────────────────────────────────────
# Style application
# ────────────────────────────────────────────────────────────────────

def _apply_fonts_and_sizes(doc, lock: Dict[str, Any]) -> None:
    """Apply font / size / line-spacing to Normal, Heading 1-3, Title, Caption."""
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    fonts = lock.get("fonts", {}) or {}
    cjk = fonts.get("cjk", "標楷體")
    latin = fonts.get("latin", "Times New Roman")
    fmt = lock.get("formatting", {}) or {}
    line_spacing = float(lock.get("line_spacing",
                                  fmt.get("body", {}).get("line_spacing", 1.5)))

    def _set_rfonts(rPr, latin_name: str, cjk_name: str) -> None:
        existing = rPr.find(qn("w:rFonts"))
        if existing is not None:
            rPr.remove(existing)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), latin_name)
        rFonts.set(qn("w:hAnsi"), latin_name)
        rFonts.set(qn("w:eastAsia"), cjk_name)
        rFonts.set(qn("w:cs"), latin_name)
        rFonts.set(qn("w:hint"), "eastAsia")
        rPr.insert(0, rFonts)

    def _set_size(rPr, size_pt: int) -> None:
        existing = rPr.find(qn("w:sz"))
        if existing is not None:
            rPr.remove(existing)
        existing_cs = rPr.find(qn("w:szCs"))
        if existing_cs is not None:
            rPr.remove(existing_cs)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))  # half-points
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(int(size_pt * 2)))
        rPr.append(szCs)

    def _set_line_spacing(pPr, multiplier: float) -> None:
        existing = pPr.find(qn("w:spacing"))
        if existing is not None:
            pPr.remove(existing)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), str(int(240 * multiplier)))
        spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)

    def _set_bold(rPr, bold: bool) -> None:
        for tag in ("w:b", "w:bCs"):
            existing = rPr.find(qn(tag))
            if existing is not None:
                rPr.remove(existing)
        if bold:
            b = OxmlElement("w:b")
            rPr.append(b)
            bCs = OxmlElement("w:bCs")
            rPr.append(bCs)

    def _style_apply(style, size_pt: Optional[int], bold: bool = False) -> None:
        el = style.element
        rPr = el.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            el.append(rPr)
        _set_rfonts(rPr, latin, cjk)
        if size_pt is not None:
            _set_size(rPr, size_pt)
        if bold:
            _set_bold(rPr, True)
        # Line spacing at paragraph level
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            rPr_el = el.find(qn("w:rPr"))
            if rPr_el is not None:
                rPr_el.addprevious(pPr)
            else:
                el.append(pPr)
        _set_line_spacing(pPr, line_spacing)

    # Normal (body)
    body_fmt = fmt.get("body", {}) or {}
    _style_apply(doc.styles["Normal"],
                 int(body_fmt.get("font_size", 12)),
                 bold=bool(body_fmt.get("bold", False)))

    # Headings 1-3
    for level in (1, 2, 3):
        hkey = f"h{level}"
        hfmt = fmt.get(hkey, {}) or {}
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        _style_apply(style, int(hfmt.get("font_size", {1: 18, 2: 16, 3: 14}[level])),
                     bold=bool(hfmt.get("bold", True)))

    # Title
    if "Title" in [s.name for s in doc.styles]:
        t_fmt = fmt.get("title", {}) or {}
        try:
            _style_apply(doc.styles["Title"],
                         int(t_fmt.get("font_size", 22)),
                         bold=bool(t_fmt.get("bold", True)))
        except KeyError:
            pass


def _apply_page_setup(doc, lock: Dict[str, Any]) -> None:
    """Apply page size + margins from lock to all sections."""
    from docx.shared import Cm

    page_size_name = str(lock.get("page_size", "A4")).strip()
    section = doc.sections[0]
    if page_size_name in _PAGE_SIZES_CM:
        w_cm, h_cm = _PAGE_SIZES_CM[page_size_name]
        section.page_width = Cm(w_cm)
        section.page_height = Cm(h_cm)
    margins = lock.get("margins", {}) or {}
    try:
        section.top_margin = Cm(_to_cm(margins.get("top"), 2.5))
        section.bottom_margin = Cm(_to_cm(margins.get("bottom"), 2.5))
        section.left_margin = Cm(_to_cm(margins.get("left"), 3.0))
        section.right_margin = Cm(_to_cm(margins.get("right"), 2.0))
    except Exception as e:
        logger.warning("margin application failed: %s", e)


# ────────────────────────────────────────────────────────────────────
# HTML parsing → docx
# ────────────────────────────────────────────────────────────────────

def _is_block(tag_name: str) -> bool:
    return tag_name in {
        "p", "div", "section", "article", "header", "footer",
        "h1", "h2", "h3", "h4", "h5", "h6",
        "ul", "ol", "li",
        "table", "tr", "td", "th", "thead", "tbody", "tfoot",
        "blockquote", "pre",
        "hr", "br",
    }


def _get_inline_text(node) -> str:
    """Recursively extract text from a node, preserving basic whitespace."""
    from bs4 import NavigableString
    if isinstance(node, NavigableString):
        return str(node)
    return node.get_text()


def _sanitize_markdown_emphasis(html: str) -> str:
    """Convert stray Markdown bold/italic syntax to HTML tags.

    Why: upstream producers (e.g. LLM-generated section HTML) sometimes
    emit literal ``**text**`` instead of ``<strong>text</strong>``. If we
    pass that string straight to python-docx, the asterisks survive as
    literal characters and Word displays ``**粗體**`` instead of bold text.

    This sanitizer runs on the raw HTML string before BeautifulSoup parses
    it, so downstream ``handle_strong``/``handle_b`` always see clean tags.

    Rules (conservative — only operates on text NOT already inside an HTML
    tag, to avoid corrupting attribute values):

    - ``**text**``     → ``<strong>text</strong>``   (Markdown bold)
    - ``__text__``     → ``<strong>text</strong>``   (Markdown bold alt)
    - ``*text*``       → ``<em>text</em>``           (Markdown italic)
    - ``_text_``       → ``<em>text</em>``           (Markdown italic alt)

    Non-greedy matching + minimum 1 char inside prevents matching empty
    pairs. Whitespace-only content is left alone.
    """
    if not html or "**" not in html and "__" not in html and "*" not in html and "_" not in html:
        return html
    # Use a callback so we never match across HTML tag boundaries.
    out_parts: List[str] = []
    i = 0
    n = len(html)
    while i < n:
        ch = html[i]
        if ch == "<":
            # Copy the entire tag verbatim until the matching '>'.
            end = html.find(">", i)
            if end == -1:
                out_parts.append(html[i:])
                break
            out_parts.append(html[i : end + 1])
            i = end + 1
            continue
        if ch == "*":
            # Try **bold**
            if i + 1 < n and html[i + 1] == "*":
                end = html.find("**", i + 2)
                if end != -1 and end > i + 2:
                    inner = html[i + 2 : end]
                    if inner.strip():
                        out_parts.append(f"<strong>{inner}</strong>")
                        i = end + 2
                        continue
            # Try *italic*
            end = html.find("*", i + 1)
            if end != -1 and end > i + 1:
                inner = html[i + 1 : end]
                if inner.strip() and "<" not in inner and ">" not in inner:
                    out_parts.append(f"<em>{inner}</em>")
                    i = end + 1
                    continue
            out_parts.append(ch)
            i += 1
            continue
        if ch == "_":
            # Try __bold__
            if i + 1 < n and html[i + 1] == "_":
                end = html.find("__", i + 2)
                if end != -1 and end > i + 2:
                    inner = html[i + 2 : end]
                    if inner.strip():
                        out_parts.append(f"<strong>{inner}</strong>")
                        i = end + 2
                        continue
            # Try _italic_
            end = html.find("_", i + 1)
            if end != -1 and end > i + 1:
                inner = html[i + 1 : end]
                if inner.strip() and "<" not in inner and ">" not in inner:
                    out_parts.append(f"<em>{inner}</em>")
                    i = end + 1
                    continue
            out_parts.append(ch)
            i += 1
            continue
        out_parts.append(ch)
        i += 1
    return "".join(out_parts)


def _add_runs(paragraph, node, lock: Dict[str, Any], base_bold: bool = False) -> None:
    """Recursively add runs to a paragraph from an HTML node.

    Handles <strong>/<b>, <em>/<i>, <code>, <u>, <a>, <sup>, <sub>.

    Bold handling contract (Problem 4 fix):
        - When the node is ``<strong>`` or ``<b>``, every descendant run
          is created with ``run.bold = True``.
        - When the node already contains another ``<strong>``/``<b>``,
          we OR the bold flag with the existing one (idempotent).
        - When the input HTML still contains literal ``**`` markers
          (escape failure upstream), ``_sanitize_markdown_emphasis()``
          converts them to ``<strong>`` before we ever see the node.
    """
    from bs4 import NavigableString, Tag
    from docx.shared import RGBColor

    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        if base_bold:
            run.bold = True
        return

    if not isinstance(node, Tag):
        return

    name = (node.name or "").lower()
    children = list(node.children)

    if name in ("strong", "b"):
        # Recurse with base_bold OR'd in, so nested <strong><strong>x
        # </strong></strong> still produces one bold run (idempotent).
        for child in children:
            _add_runs(paragraph, child, lock, base_bold=base_bold or True)
        return
    if name in ("em", "i"):
        for child in children:
            _add_runs(paragraph, child, lock, base_bold=base_bold)
            # mark last run italic
            if paragraph.runs:
                paragraph.runs[-1].italic = True
        return
    if name == "u":
        for child in children:
            _add_runs(paragraph, child, lock, base_bold=base_bold)
            if paragraph.runs:
                paragraph.runs[-1].underline = True
        return
    if name == "code":
        text = node.get_text()
        run = paragraph.add_run(text)
        run.font.name = "Courier New"
        # Tag eastAsia for consistency
        return
    if name == "a":
        text = node.get_text()
        href = node.get("href", "")
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        run.underline = True
        if href and not href.startswith("#"):
            run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        return
    if name == "sup":
        text = node.get_text()
        run = paragraph.add_run(text)
        run.font.superscript = True
        return
    if name == "sub":
        text = node.get_text()
        run = paragraph.add_run(text)
        run.font.subscript = True
        return
    if name == "img":
        # Inline image inside a paragraph. Delegate to the image helper.
        try:
            from docx.shared import Inches
            import urllib.parse
            import urllib.request
            import tempfile
            from io import BytesIO
            import base64 as _b64
            src = node.get("src", "").strip()
            if not src:
                return
            img_source = None
            cleanup = None
            if src.startswith("data:"):
                head, b64 = src.split(",", 1)
                data = _b64.b64decode(b64)
                img_source = BytesIO(data)
            elif src.startswith(("http://", "https://")):
                with urllib.request.urlopen(src) as resp:  # nosec
                    tmp = tempfile.NamedTemporaryFile(suffix=".img", delete=False)
                    tmp.write(resp.read())
                    tmp.close()
                    img_source = Path(tmp.name)
                    cleanup = img_source
            else:
                if src.startswith("file://"):
                    src = urllib.parse.unquote(src[len("file://"):])
                p_path = Path(src)
                if not p_path.is_absolute():
                    p_path = (Path.cwd() / p_path).resolve()
                if p_path.exists():
                    img_source = p_path
                else:
                    paragraph.add_run(f"[image: {src}]")
                    return
            width_attr = node.get("width")
            try:
                w_emu = Inches(float(width_attr.rstrip("px")) / 96) if width_attr else Inches(5)
            except Exception:
                w_emu = Inches(5)
            run = paragraph.add_run()
            if hasattr(img_source, "read"):
                run.add_picture(img_source, width=w_emu)
            else:
                run.add_picture(str(img_source), width=w_emu)
            if cleanup is not None:
                try:
                    cleanup.unlink()
                except OSError:
                    pass
        except Exception as e:
            paragraph.add_run(f"[image: {node.get('alt') or node.get('src','')}]")
            logger.warning("inline image failed: %s", e)
        return
    if name == "br":
        paragraph.add_run().add_break()
        return

    # Default: recurse into children
    for child in children:
        _add_runs(paragraph, child, lock, base_bold=base_bold)


def _add_paragraph(doc, node, lock: Dict[str, Any]) -> None:
    """Add a <p>-like block as a paragraph to doc."""
    p = doc.add_paragraph(style="Normal")
    _add_runs(p, node, lock)


def _add_heading(doc, node, level: int) -> None:
    """Add a heading. level is 1..6 (clamped to 1..3 to match available styles)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    level = max(1, min(level, 3))  # we only customize 1..3
    style_name = f"Heading {level}"
    text = node.get_text()
    p = doc.add_paragraph(text, style=style_name)
    return p


def _add_list(doc, node, ordered: bool, lock: Dict[str, Any]) -> None:
    """Add a <ul> or <ol> with its <li> children.

    Uses Word's real List Number / List Bullet style so Word treats them
    as structured lists (auto-renumber on delete, proper list indentation).
    Falls back to manual prefix + Normal style if the style doesn't exist.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Determine which style to use
    style_name = "List Number" if ordered else "List Bullet"
    style_names = [s.name for s in doc.styles]
    use_real_list = style_name in style_names

    items = [c for c in node.children
             if getattr(c, "name", None) == "li"]

    for idx, li in enumerate(items, start=1):
        if use_real_list:
            p = doc.add_paragraph(style=style_name)
        else:
            # Fallback: manual prefix (old behaviour)
            marker = f"{idx}." if ordered else "•"
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(f"{marker}\t")
        # Inline content
        _add_runs(p, li, lock)
        # Indent
        try:
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "720")  # 0.5 inch
            pPr.append(ind)
        except Exception:
            pass


def _add_table(doc, node, lock: Dict[str, Any]) -> None:
    """Convert a <table> to a docx table.

    Supports <thead> / <tbody> / <tfoot>, <th> / <td>, basic colspan (visual only).
    Nested tables: outer <table> cells containing inner <table> get rendered
    with a placeholder line + ignored (python-docx cannot represent nested
    tables cleanly, so we log and skip).
    """
    rows = []
    # Collect <tr> from thead/tbody/tfoot in order, or directly from table
    sources = []
    for container_name in ("thead", "tbody", "tfoot"):
        c = node.find(container_name)
        if c:
            sources.extend(c.find_all("tr", recursive=False))
    if not sources:
        sources = node.find_all("tr", recursive=False)

    for tr in sources:
        cells = tr.find_all(["th", "td"], recursive=False)
        rows.append(cells)

    if not rows:
        return

    cols = max(len(r) for r in rows)
    if cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True

    # Read lock table formatting
    table_fmt = (lock.get("formatting") or {}).get("table") or {}
    table_font_size = table_fmt.get("font_size", 12)

    for r_idx, cells in enumerate(rows):
        for c_idx, cell_html in enumerate(cells):
            if c_idx >= cols:
                break
            cell = table.rows[r_idx].cells[c_idx]
            # Clear default empty paragraph
            for p in list(cell.paragraphs):
                p._element.getparent().remove(p._element)
            p = cell.add_paragraph(style="Normal")
            # Bold if <th>
            is_header = (cell_html.name == "th")
            _add_runs(p, cell_html, lock, base_bold=is_header)
            # Apply table font size to all runs in this cell
            for run in p.runs:
                run.font.size = Pt(table_font_size)
            # If cell contains a nested <table>, log and continue
            if cell_html.find("table") != -1:
                logger.info("nested <table> detected; docx can't represent it cleanly")


def _add_image(doc, node, lock: Dict[str, Any]) -> None:
    """Add an <img> as an inline image in a new paragraph.

    Supports: src as file path / file:// / data: URI / http(s)://.
    For http(s), the image is downloaded to a temp file first.
    Aspect ratio is preserved by only setting width (height auto).
    """
    import urllib.parse
    import urllib.request
    import tempfile
    from docx.shared import Inches

    src = node.get("src", "").strip()
    if not src:
        return
    alt = node.get("alt", "")

    p = doc.add_paragraph()
    p.alignment = 1  # center
    run = p.add_run()
    img_source: Optional[Union[str, Path]] = None
    cleanup_path: Optional[Path] = None
    try:
        if src.startswith("data:"):
            # data:image/png;base64,... → BytesIO
            from io import BytesIO
            import base64
            head, b64 = src.split(",", 1)
            data = base64.b64decode(b64)
            img_source = BytesIO(data)
        elif src.startswith(("http://", "https://")):
            with urllib.request.urlopen(src) as resp:  # nosec - best effort
                data = resp.read()
            tmp = tempfile.NamedTemporaryFile(suffix=".img", delete=False)
            tmp.write(data)
            tmp.close()
            img_source = Path(tmp.name)
            cleanup_path = img_source
        else:
            # Local path (handle file:// too)
            if src.startswith("file://"):
                src = urllib.parse.unquote(src[len("file://"):])
            p_path = Path(src)
            if not p_path.is_absolute():
                p_path = (Path.cwd() / p_path).resolve()
            if p_path.exists():
                img_source = p_path
            else:
                logger.warning("image not found: %s", p_path)
                run.add_text(f"[image: {alt or src}]")
                return

        # Default width: 5 inches (let height auto to keep aspect ratio)
        width = node.get("width")
        try:
            w_emu = Inches(float(width.rstrip("px")) / 96) if width else Inches(5)
        except Exception:
            w_emu = Inches(5)
        if hasattr(img_source, "read"):
            run.add_picture(img_source, width=w_emu)  # BytesIO
        else:
            run.add_picture(str(img_source), width=w_emu)
    except Exception as e:
        logger.warning("failed to add image %s: %s", src, e)
        run.add_text(f"[image: {alt or src}]")
    finally:
        if cleanup_path is not None:
            try:
                cleanup_path.unlink()
            except OSError:
                pass


def _add_blockquote(doc, node, lock: Dict[str, Any]) -> None:
    """Add a <blockquote> as an indented italic paragraph."""
    p = doc.add_paragraph(style="Normal")
    _add_runs(p, node, lock)
    for run in p.runs:
        run.italic = True
    # Indent
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:right"), "720")
        pPr.append(ind)
    except Exception:
        pass


def _add_pre(doc, node, lock: Dict[str, Any]) -> None:
    """Add a <pre> as a monospace paragraph (preserve newlines).

    Applies lock["formatting"]["code"] font_name / font_size if available,
    falling back to Courier New 10pt.
    """
    from docx.shared import Pt

    code_fmt = (lock.get("formatting") or {}).get("code") or {}
    font_name = code_fmt.get("font_name") or code_fmt.get("font") or "Courier New"
    font_size = code_fmt.get("font_size", 10)

    text = node.get_text()
    for line in text.splitlines():
        p = doc.add_paragraph(style="Normal")
        run = p.add_run(line if line else " ")
        run.font.name = font_name
        run.font.size = Pt(font_size)


def _add_hr(doc) -> None:
    """Add a horizontal rule (paragraph with bottom border).

    NOTE: uses w:pBdr bottom border — may conflict with paragraph spacing
    on narrow pages. Long-term fix: use set_border() API or w:sectPr.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _walk_node(doc, node, lock: Dict[str, Any]) -> None:
    """Recursively walk an HTML node tree and emit docx blocks."""
    from bs4 import NavigableString, Tag

    if isinstance(node, NavigableString):
        # Stray text outside a block — wrap in a paragraph
        text = str(node).strip()
        if text:
            p = doc.add_paragraph(text, style="Normal")
        return

    if not isinstance(node, Tag):
        return

    name = (node.name or "").lower()

    # Skip <head>, <style>, <script>, <meta>, <link>
    if name in ("head", "style", "script", "meta", "link", "title", "noscript"):
        return

    # Headings
    m = re.match(r"^h([1-6])$", name)
    if m:
        _add_heading(doc, node, int(m.group(1)))
        return

    if name == "p":
        _add_paragraph(doc, node, lock)
        return
    if name == "ul":
        _add_list(doc, node, ordered=False, lock=lock)
        return
    if name == "ol":
        _add_list(doc, node, ordered=True, lock=lock)
        return
    if name == "table":
        _add_table(doc, node, lock)
        return
    if name == "img":
        _add_image(doc, node, lock)
        return
    if name == "blockquote":
        _add_blockquote(doc, node, lock)
        return
    if name == "pre":
        _add_pre(doc, node, lock)
        return
    if name == "hr":
        _add_hr(doc)
        return
    if name == "br":
        # Lone <br> outside a paragraph: just a blank line
        doc.add_paragraph("", style="Normal")
        return

    # Generic container: recurse into children
    for child in node.children:
        _walk_node(doc, child, lock)


# ────────────────────────────────────────────────────────────────────
# Public API
# ────────────────────────────────────────────────────────────────────

def html_to_docx_direct(
    html_source: Union[str, Path],
    output_docx: Union[str, Path],
    *,
    template: Optional[Union[str, Path]] = None,
    lock_file: Optional[Union[str, Path]] = None,
    cjk_font: Optional[str] = None,
    latin_font: Optional[str] = None,
) -> Path:
    """Convert HTML → DOCX using python-docx (high-control path).

    Args:
        html_source: HTML string OR path to HTML file.
        output_docx: path to write DOCX.
        template: optional path to a reference docx (loaded as base).
        lock_file: path to report_lock.md; defaults to examples/lock.md.
        cjk_font / latin_font: explicit font overrides (bypass lock).

    Returns: Path to output_docx.

    Raises:
        HTMLToDOCXDirectError: on failure.
    """
    _require_deps()
    from bs4 import BeautifulSoup
    from docx import Document

    out = Path(output_docx)
    out.parent.mkdir(parents=True, exist_ok=True)

    # ─── Load lock ───
    if lock_file is None:
        lock_path = _PROJECT_ROOT / "examples" / "lock.md"
    else:
        lock_path = Path(lock_file)
    lock = _load_lock(lock_path)

    # Allow explicit overrides
    if cjk_font or latin_font:
        fonts = dict(lock.get("fonts", {}))
        if cjk_font:
            fonts["cjk"] = cjk_font
        if latin_font:
            fonts["latin"] = latin_font
        lock["fonts"] = fonts

    # ─── Load HTML ───
    raw_html: Optional[str] = None
    if isinstance(html_source, Path):
        if not html_source.exists():
            raise HTMLToDOCXDirectError(f"HTML 檔不存在: {html_source}")
        if html_source.stat().st_size == 0:
            raise HTMLToDOCXDirectError(f"HTML 檔為空: {html_source}")
        raw_html = html_source.read_text(encoding="utf-8")
    else:
        s = str(html_source or "")
        s_stripped = s.strip()
        # Decide: path-like string vs. inline HTML
        looks_like_path = (
            s_stripped
            and "\n" not in s
            and "<" not in s
            and (
                s_stripped.endswith((".html", ".htm", ".xhtml"))
                or s_stripped.startswith(("/", "./", "../", "~"))
            )
        )
        if looks_like_path:
            p = Path(s_stripped)
            if not p.exists():
                raise HTMLToDOCXDirectError(f"HTML 檔不存在: {p}")
            if p.stat().st_size == 0:
                raise HTMLToDOCXDirectError(f"HTML 檔為空: {p}")
            raw_html = p.read_text(encoding="utf-8")
        else:
            if not s_stripped:
                raise HTMLToDOCXDirectError("HTML 字串為空")
            raw_html = s

    # ─── Sanitize stray Markdown bold/italic syntax ───
    # Upstream producers sometimes emit literal `**text**` instead of
    # `<strong>text</strong>`. Convert those into proper tags before
    # parsing, so handle_strong/handle_b see clean markup and apply
    # `bold=True` instead of leaving the asterisks as literal chars.
    raw_html = _sanitize_markdown_emphasis(raw_html)

    # Parse (silence the spurious 'looks like a filename' warning that
    # lxml emits when the HTML doesn't begin with a recognized doctype/tag).
    import warnings
    from bs4 import MarkupResemblesLocatorWarning
    warnings.filterwarnings("ignore", category=MarkupResemblesLocatorWarning)
    soup = BeautifulSoup(raw_html, "lxml")
    body = soup.body or soup

    # ─── Build document ───
    if template is not None:
        tpl = Path(template)
        if not tpl.exists():
            raise HTMLToDOCXDirectError(f"template 不存在: {tpl}")
        doc = Document(str(tpl))
    else:
        doc = Document()

    _apply_fonts_and_sizes(doc, lock)
    _apply_page_setup(doc, lock)

    # Walk DOM
    for child in list(body.children):
        _walk_node(doc, child, lock)

    # ─── Save ───
    doc.save(str(out))

    if not out.exists() or out.stat().st_size == 0:
        raise HTMLToDOCXDirectError(f"DOCX 寫入失敗或為空: {out}")

    logger.info("DOCX (python-docx) written: %s (%d bytes)",
                out, out.stat().st_size)
    return out


# ────────────────────────────────────────────────────────────────────
# CLI
# ────────────────────────────────────────────────────────────────────

def _main(argv=None) -> int:
    ap = argparse.ArgumentParser(
        prog="scripts.html_to_docx_direct",
        description=(
            "Convert HTML → DOCX via python-docx (parallel path, "
            "bypasses Markdown)."
        ),
    )
    ap.add_argument("--input", "-i", required=True,
                    help="Input HTML file path")
    ap.add_argument("--output", "-o", required=True,
                    help="Output DOCX path")
    ap.add_argument("--template", default=None,
                    help="Reference docx template (default: blank document)")
    ap.add_argument("--lock-file", default=None,
                    help="Path to report_lock.md (default: examples/lock.md)")
    ap.add_argument("--cjk-font", default=None,
                    help="Override CJK font (bypasses lock)")
    ap.add_argument("--latin-font", default=None,
                    help="Override Latin font (bypasses lock)")
    ap.add_argument("--verbose", "-v", action="store_true",
                    help="Verbose logging")
    args = ap.parse_args(argv)

    if args.verbose:
        logging.basicConfig(level=logging.INFO,
                            format="%(levelname)s %(name)s: %(message)s")

    try:
        out = html_to_docx_direct(
            html_source=args.input,
            output_docx=args.output,
            template=args.template,
            lock_file=args.lock_file,
            cjk_font=args.cjk_font,
            latin_font=args.latin_font,
        )
        size = out.stat().st_size
        print(f"✅ DOCX written: {out} ({size} bytes)")
        return 0
    except HTMLToDOCXDirectError as e:
        print(f"❌ {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(_main())
