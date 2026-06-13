"""tests/test_source_to_md.py — source_to_md/* 單元測試。

對應 `tasks.md` T1-3 / T1-4 / T1-5 / T1-6 + DoD：
- 給 1 個簡單 PDF / DOCX 跑 normalize
- md_normalizer 統一測試
- 缺檔 / 格式錯誤情境

注意：PDF / DOCX 產生使用 PyMuPDF / python-docx（python-docx 不在 requirements，
但用來建測試 fixture）。
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

import pytest

# 測試 fixtures 用的 optional import
try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from scripts.source_to_md.md_normalizer import (
    extract_frontmatter,
    has_frontmatter,
    normalize,
    normalize_file,
    write_normalized,
)


# ─── md_normalizer 測試 ──────────────────────────────────────────────

class TestMdNormalizer:
    """md_normalizer.py 核心測試。"""

    def test_strips_utf8_bom(self):
        """應移除 UTF-8 BOM。"""
        text = "\ufeff# 標題\n內文"
        out = normalize(text)
        assert not out.startswith("\ufeff")
        assert out.startswith("# 標題")

    def test_normalizes_crlf_to_lf(self):
        """CRLF / CR → LF。"""
        text = "# 標題\r\n\r\n內文\r\n"
        out = normalize(text)
        assert "\r" not in out
        assert out.count("\n") >= 3

    def test_strips_trailing_whitespace(self):
        """每行不應有 trailing whitespace。"""
        text = "# 標題   \n內文\t\n\n"
        out = normalize(text)
        for line in out.split("\n"):
            assert line == line.rstrip(), f"trailing ws in: {line!r}"

    def test_collapses_multiple_blank_lines(self):
        """3+ 連續空行 → 最多 2（即 1 個段落分隔）。"""
        text = "段落一\n\n\n\n\n段落二"
        out = normalize(text)
        assert "\n\n\n" not in out

    def test_normalizes_heading_prefix(self):
        """標題前綴統一單一空白。"""
        text = "##  H2\n###   H3\n####    H4"
        out = normalize(text)
        for line in out.split("\n"):
            if line.startswith("#"):
                # 計算 # 後的空白
                hash_count = len(line) - len(line.lstrip("#"))
                rest = line[hash_count:]
                assert rest.startswith(" "), f"heading without space: {line!r}"
                assert not rest.startswith("  "), f"heading with extra space: {line!r}"

    def test_frontmatter_preserved(self):
        """有 frontmatter 時應保留。"""
        text = "---\ntitle: Test\nauthor: Zero\n---\n# 標題\n內文"
        out = normalize(text)
        assert "---" in out
        assert "title: Test" in out
        assert "# 標題" in out

    def test_normalize_file(self, tmp_path: Path):
        """normalize_file 應讀檔並回傳。"""
        p = tmp_path / "in.md"
        p.write_text("\ufeff# 標題\r\n\r\n內文   \r\n", encoding="utf-8")
        out = normalize_file(p)
        assert "標題" in out
        assert "\r" not in out
        assert "   " not in out

    def test_write_normalized(self, tmp_path: Path):
        """write_normalized 應寫檔。"""
        p = tmp_path / "out.md"
        write_normalized(p, "\ufeff# 標題\r\n內文   \r\n")
        assert p.exists()
        content = p.read_text(encoding="utf-8")
        assert "\ufeff" not in content
        assert "\r" not in content

    def test_extract_frontmatter(self):
        """應正確解析 frontmatter。"""
        text = "---\ntitle: X\ncount: 3\n---\n# Body\n內文"
        data, body = extract_frontmatter(text)
        assert data is not None
        assert data["title"] == "X"
        assert data["count"] == 3
        assert "# Body" in body

    def test_extract_frontmatter_no_fm(self):
        """無 frontmatter 時回 (None, 原文)。"""
        text = "# Body\n內文"
        data, body = extract_frontmatter(text)
        assert data is None
        assert body == text

    def test_has_frontmatter(self):
        """has_frontmatter 正確判斷。"""
        assert has_frontmatter("---\nx: 1\n---\nbody")
        assert not has_frontmatter("# Just body\n")


# ─── pdf_to_md 測試 ─────────────────────────────────────────────────

@pytest.mark.skipif(not HAS_FITZ, reason="PyMuPDF not installed")
class TestPdfToMd:
    """pdf_to_md.py 測試（建立最小 PDF fixture）。"""

    @pytest.fixture
    def sample_pdf(self, tmp_path: Path) -> Path:
        """建立一個簡單的 PDF（2 頁，含 1 個標題、1 個段落、1 張圖）。

        註：PyMuPDF 預設 helv 字體不支援 CJK，故測試用 ASCII 文字。
        邏輯（heading 推斷）不依賴語言。
        """
        pdf_path = tmp_path / "sample.pdf"
        doc = fitz.open()
        try:
            # 第 1 頁：標題 + 內文
            page1 = doc.new_page()
            page1.insert_text(
                (50, 50), "Chapter 1 Introduction",
                fontsize=18, fontname="helv",
            )
            page1.insert_text(
                (50, 100), "This is the background of the research.",
                fontsize=12, fontname="helv",
            )
            page1.insert_text(
                (50, 150), "1.1 Motivation",
                fontsize=14, fontname="helv",
            )
            page1.insert_text(
                (50, 200), "Explain the importance.",
                fontsize=12, fontname="helv",
            )

            # 第 2 頁：另一章
            page2 = doc.new_page()
            page2.insert_text(
                (50, 50), "Chapter 2 Literature Review",
                fontsize=18, fontname="helv",
            )
            page2.insert_text(
                (50, 100), "Related work.",
                fontsize=12, fontname="helv",
            )

            doc.save(str(pdf_path))
        finally:
            doc.close()
        return pdf_path

    def test_convert_creates_md(self, sample_pdf: Path, tmp_path: Path):
        """應產出 .md 檔。"""
        from scripts.source_to_md.pdf_to_md import convert
        out = tmp_path / "out" / "sample.md"
        result = convert(sample_pdf, out)
        assert result.exists()
        assert result.suffix == ".md"

    def test_convert_preserves_headings(self, sample_pdf: Path, tmp_path: Path):
        """標題層級應保留（字體大小對應 H1/H2/H3）。"""
        from scripts.source_to_md.pdf_to_md import convert
        out = tmp_path / "out.md"
        convert(sample_pdf, out)
        content = out.read_text(encoding="utf-8")
        # 第一章 + 第二章 → 兩個 H1
        assert content.count("# Chapter 1") >= 1
        assert content.count("# Chapter 2") >= 1
        # 1.1 → H3
        assert "### 1.1" in content or "### 1.1 Motivation" in content

    def test_convert_creates_frontmatter(self, sample_pdf: Path, tmp_path: Path):
        """應產 frontmatter。"""
        from scripts.source_to_md.pdf_to_md import convert
        out = tmp_path / "out.md"
        convert(sample_pdf, out)
        content = out.read_text(encoding="utf-8")
        assert content.startswith("---\n")
        assert "source_type: pdf" in content
        assert "sample.pdf" in content
        assert "page_count: 2" in content

    def test_convert_output_is_normalized(
        self, sample_pdf: Path, tmp_path: Path
    ):
        """輸出應經 md_normalizer 處理（無 CRLF、無 BOM）。"""
        from scripts.source_to_md.pdf_to_md import convert
        out = tmp_path / "out.md"
        convert(sample_pdf, out)
        content = out.read_text(encoding="utf-8")
        assert "\ufeff" not in content
        assert "\r" not in content

    def test_convert_missing_file_raises(self, tmp_path: Path):
        """找不到 PDF 應 raise FileNotFoundError。"""
        from scripts.source_to_md.pdf_to_md import convert
        with pytest.raises(FileNotFoundError):
            convert(tmp_path / "nonexistent.pdf", tmp_path / "out.md")


# ─── docx_to_md 測試 ─────────────────────────────────────────────────

@pytest.mark.skipif(
    not (HAS_MAMMOTH and HAS_DOCX),
    reason="mammoth or python-docx not installed",
)
class TestDocxToMd:
    """docx_to_md.py 測試（建立最小 DOCX fixture）。"""

    @pytest.fixture
    def sample_docx(self, tmp_path: Path) -> Path:
        """建立一個簡單的 DOCX（標題、段落、表格）。"""
        docx_path = tmp_path / "sample.docx"
        doc = docx.Document()
        try:
            doc.add_heading("第一章 緒論", level=1)
            doc.add_paragraph("這是內文段落，包含 **粗體** 與 *斜體*。")
            doc.add_heading("1.1 研究背景", level=2)
            doc.add_paragraph("說明背景。")

            # 表格
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "項目"
            table.cell(0, 1).text = "數值"
            table.cell(1, 0).text = "字數"
            table.cell(1, 1).text = "12,345"

            doc.save(str(docx_path))
        finally:
            pass
        return docx_path

    def test_convert_creates_md(self, sample_docx: Path, tmp_path: Path):
        """應產出 .md 檔。"""
        from scripts.source_to_md.docx_to_md import convert
        out = tmp_path / "out" / "sample.md"
        result = convert(sample_docx, out)
        assert result.exists()

    def test_convert_preserves_headings(self, sample_docx: Path, tmp_path: Path):
        """標題層級應保留。"""
        from scripts.source_to_md.docx_to_md import convert
        out = tmp_path / "out.md"
        convert(sample_docx, out)
        content = out.read_text(encoding="utf-8")
        assert "# 第一章 緒論" in content
        assert "## 1.1 研究背景" in content

    def test_convert_preserves_bold_italic(
        self, sample_docx: Path, tmp_path: Path
    ):
        """粗體與斜體應保留。"""
        from scripts.source_to_md.docx_to_md import convert
        out = tmp_path / "out.md"
        convert(sample_docx, out)
        content = out.read_text(encoding="utf-8")
        assert "**粗體**" in content
        assert "*斜體*" in content

    def test_convert_preserves_table(self, sample_docx: Path, tmp_path: Path):
        """表格應保留為 Markdown table。"""
        from scripts.source_to_md.docx_to_md import convert
        out = tmp_path / "out.md"
        convert(sample_docx, out)
        content = out.read_text(encoding="utf-8")
        # 應有 | 表格分隔
        assert "|" in content
        assert "項目" in content
        assert "數值" in content
        assert "字數" in content
        assert "12,345" in content

    def test_convert_creates_frontmatter(
        self, sample_docx: Path, tmp_path: Path
    ):
        """應產 frontmatter。"""
        from scripts.source_to_md.docx_to_md import convert
        out = tmp_path / "out.md"
        convert(sample_docx, out)
        content = out.read_text(encoding="utf-8")
        assert "source_type: docx" in content
        assert "sample.docx" in content

    def test_convert_output_is_normalized(
        self, sample_docx: Path, tmp_path: Path
    ):
        """輸出應經 normalize 處理。"""
        from scripts.source_to_md.docx_to_md import convert
        out = tmp_path / "out.md"
        convert(sample_docx, out)
        content = out.read_text(encoding="utf-8")
        assert "\ufeff" not in content
        assert "\r" not in content

    def test_convert_missing_file_raises(self, tmp_path: Path):
        """找不到 DOCX 應 raise FileNotFoundError。"""
        from scripts.source_to_md.docx_to_md import convert
        with pytest.raises(FileNotFoundError):
            convert(tmp_path / "nonexistent.docx", tmp_path / "out.md")


# ─── url_to_md 測試 ─────────────────────────────────────────────────

class TestUrlToMd:
    """url_to_md.py 測試（不真實抓網路，用 mock）。"""

    def test_convert_with_mock_html(self, tmp_path: Path, monkeypatch):
        """用 mock HTML 測試 converter。"""
        from scripts.source_to_md import url_to_md

        sample_html = """<!DOCTYPE html>
<html lang="zh-TW">
<head>
  <title>範例文章</title>
</head>
<body>
  <nav><a href="/">首頁</a></nav>
  <article>
    <h1>標題一</h1>
    <p>這是第一段，包含 <strong>粗體</strong> 與 <em>斜體</em>。</p>
    <h2>小標題</h2>
    <ul>
      <li>項目 A</li>
      <li>項目 B</li>
    </ul>
    <p>參考 <a href="https://example.com">連結</a>。</p>
  </article>
  <script>alert('x')</script>
  <style>body { color: red; }</style>
</body>
</html>
"""

        class _MockResp:
            text = sample_html
            encoding = "utf-8"
            apparent_encoding = "utf-8"

            def raise_for_status(self):
                pass

        def _mock_get(url, headers=None, timeout=None):
            return _MockResp()

        monkeypatch.setattr(url_to_md.requests, "get", _mock_get)

        out = tmp_path / "out.md"
        result = url_to_md.convert("https://example.com/article", out)
        assert result.exists()
        content = result.read_text(encoding="utf-8")

        # frontmatter
        assert "source_type: url" in content
        assert "https://example.com/article" in content
        assert "title:" in content and "範例文章" in content

        # 標題
        assert "# 標題一" in content
        assert "## 小標題" in content

        # 粗體斜體
        assert "**粗體**" in content
        assert "*斜體*" in content

        # 列表
        assert "- 項目 A" in content
        assert "- 項目 B" in content

        # 連結
        assert "[連結](https://example.com)" in content

        # script / style 應被跳過（不出現在輸出或至少不出現在 body）
        assert "alert" not in content
        assert "color: red" not in content

    def test_convert_missing_file_path(self, tmp_path: Path, monkeypatch):
        """測試網路錯誤時的處理。"""
        from scripts.source_to_md import url_to_md

        def _mock_get_fail(url, headers=None, timeout=None):
            raise url_to_md.requests.RequestException("404 Not Found")

        monkeypatch.setattr(url_to_md.requests, "get", _mock_get_fail)

        out = tmp_path / "out.md"
        with pytest.raises(url_to_md.requests.RequestException):
            url_to_md.convert("https://nonexistent.example.com", out)
