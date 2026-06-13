# test_html_to_docx_direct.py — tests for scripts.html_to_docx_direct
# 對應 tasks.md T2-3
# 涵蓋：
#   - import / CLI 不 crash
#   - examples/output_1 與 output_2 都能成功轉成 .docx 且 > 1KB
#   - <img> 元素會被轉成 inline image
#   - <table> 元素會被轉成 docx table
#   - lock 設定（字體名、間距、頁面大小）有正確寫入 styles.xml

from __future__ import annotations

import base64
import io
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

import pytest
from PIL import Image

# 把 project root 加到 sys.path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from scripts.html_to_docx_direct import (
    html_to_docx_direct,
    _main,
    HTMLToDOCXDirectError,
    HTMLToDOCXDirectImportError,
    HTMLToDOCXDirectLockError,
    _require_deps,
    _to_cm,
    _apply_fonts_and_sizes,
    _apply_page_setup,
    _load_lock,
)


# ────────────────────────────────────────────────────────────────────
# Skip if required deps missing
# ────────────────────────────────────────────────────────────────────

def _has_deps() -> bool:
    try:
        _require_deps()
        return True
    except HTMLToDOCXDirectImportError:
        return False


pytestmark = pytest.mark.skipif(
    not _has_deps(),
    reason="python-docx / beautifulsoup4 / lxml not installed",
)


PROJECT_ROOT = Path(__file__).resolve().parent.parent
EXAMPLE_HTML_1 = PROJECT_ROOT / "examples" / "output_1" / "report_final.html"
EXAMPLE_HTML_2 = PROJECT_ROOT / "examples" / "output_2" / "report_final.html"
DEFAULT_LOCK = PROJECT_ROOT / "examples" / "lock.md"


# ────────────────────────────────────────────────────────────────────
# 1. Import sanity
# ────────────────────────────────────────────────────────────────────

def test_module_imports():
    """html_to_docx_direct 模組可被 import，公開 API 存在。"""
    assert callable(html_to_docx_direct)
    assert callable(_main)
    assert HTMLToDOCXDirectError is not None


def test_require_deps_passes():
    """_require_deps() 在測試環境中應該通過（套件齊備）。"""
    _require_deps()  # 不 raise


# ────────────────────────────────────────────────────────────────────
# 2. CLI smoke
# ────────────────────────────────────────────────────────────────────

def test_cli_help_does_not_crash(capsys):
    """`--help` 不 crash 並輸出 usage。"""
    with pytest.raises(SystemExit) as exc_info:
        _main(["--help"])
    # argparse 對 --help 會 exit 0
    assert exc_info.value.code == 0
    out = capsys.readouterr().out
    assert "Convert HTML" in out or "DOCX" in out


def test_cli_via_subprocess(tmp_path):
    """完整 CLI subprocess 跑 examples/output_1 → 產出 .docx > 1KB。"""
    out = tmp_path / "cli_out.docx"
    result = subprocess.run(
        [
            sys.executable, "-m", "scripts.html_to_docx_direct",
            "--input", str(EXAMPLE_HTML_1),
            "--output", str(out),
            "--lock-file", str(DEFAULT_LOCK),
        ],
        cwd=str(PROJECT_ROOT),
        capture_output=True,
        text=True,
    )
    assert result.returncode == 0, f"stderr: {result.stderr}"
    assert out.exists()
    assert out.stat().st_size > 1024


# ────────────────────────────────────────────────────────────────────
# 3. End-to-end: examples/output_1 → .docx
# ────────────────────────────────────────────────────────────────────

def test_output_1_converts_to_docx(tmp_path):
    """轉換 examples/output_1/report_final.html → 產出 .docx > 1KB。"""
    out = tmp_path / "out1.docx"
    result = html_to_docx_direct(
        html_source=EXAMPLE_HTML_1,
        output_docx=out,
        lock_file=DEFAULT_LOCK,
    )
    assert result == out
    assert out.exists()
    assert out.stat().st_size > 1024
    # Valid DOCX (ZIP)
    with zipfile.ZipFile(str(out)) as z:
        assert "[Content_Types].xml" in z.namelist()
        assert "word/document.xml" in z.namelist()


# ────────────────────────────────────────────────────────────────────
# 4. End-to-end: examples/output_2 → .docx
# ────────────────────────────────────────────────────────────────────

def test_output_2_converts_to_docx(tmp_path):
    """轉換 examples/output_2/report_final.html → 產出 .docx > 1KB。"""
    out = tmp_path / "out2.docx"
    result = html_to_docx_direct(
        html_source=EXAMPLE_HTML_2,
        output_docx=out,
        lock_file=DEFAULT_LOCK,
    )
    assert result == out
    assert out.exists()
    assert out.stat().st_size > 1024


# ────────────────────────────────────────────────────────────────────
# 5. Image conversion
# ────────────────────────────────────────────────────────────────────

def _make_test_image(color="red", size=(80, 80)) -> str:
    """Create a tiny PNG and return its data URI."""
    img = Image.new("RGB", size, color=color)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    b64 = base64.b64encode(buf.getvalue()).decode()
    return f"data:image/png;base64,{b64}"


def test_img_element_becomes_inline_image(tmp_path):
    """<img> 元素會被轉成 docx 的 inline image。"""
    data_uri = _make_test_image()
    html = f"""<!DOCTYPE html>
<html><body>
<h1>Image Test</h1>
<p><img src="{data_uri}" alt="red square" width="100"></p>
</body></html>"""
    out = tmp_path / "img.docx"
    html_to_docx_direct(html, out, lock_file=DEFAULT_LOCK)
    from docx import Document
    doc = Document(str(out))
    assert len(doc.inline_shapes) >= 1


def test_block_level_img_works(tmp_path):
    """Block-level <img>（不在 <p> 內）也能被處理。"""
    data_uri = _make_test_image(color="blue")
    html = f"""<!DOCTYPE html>
<html><body>
<h1>X</h1>
<img src="{data_uri}" alt="b" width="50">
<p>after</p>
</body></html>"""
    out = tmp_path / "blockimg.docx"
    html_to_docx_direct(html, out, lock_file=DEFAULT_LOCK)
    from docx import Document
    doc = Document(str(out))
    assert len(doc.inline_shapes) >= 1


# ────────────────────────────────────────────────────────────────────
# 6. Table conversion
# ────────────────────────────────────────────────────────────────────

def test_table_element_becomes_docx_table(tmp_path):
    """<table> 元素會被轉成 docx table，cell 文字內容保留。"""
    html = """<!DOCTYPE html>
<html><body>
<h1>Table Test</h1>
<table>
<tr><th>項目</th><th>數值</th></tr>
<tr><td>字數</td><td>12,345</td></tr>
<tr><td>頁數</td><td>42</td></tr>
</table>
</body></html>"""
    out = tmp_path / "tbl.docx"
    html_to_docx_direct(html, out, lock_file=DEFAULT_LOCK)
    from docx import Document
    doc = Document(str(out))
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert len(t.rows) == 3
    assert len(t.columns) == 2
    # First row: headers
    assert "項目" in t.rows[0].cells[0].text
    assert "數值" in t.rows[0].cells[1].text
    # Second row: data
    assert "12,345" in t.rows[1].cells[1].text


# ────────────────────────────────────────────────────────────────────
# 7. Lock settings applied (fonts, line spacing, page size)
# ────────────────────────────────────────────────────────────────────

def test_lock_fonts_applied_to_normal_style(tmp_path):
    """lock 設定的 CJK + Latin 字體應寫入 styles.xml (Normal style)。"""
    from docx import Document
    from docx.oxml.ns import qn

    html = "<html><body><h1>T</h1><p>body</p></body></html>"
    out = tmp_path / "styles.docx"
    html_to_docx_direct(html, out, lock_file=DEFAULT_LOCK)
    doc = Document(str(out))
    normal = doc.styles["Normal"]
    rpr = normal.element.find(qn("w:rPr"))
    assert rpr is not None
    rfonts = rpr.find(qn("w:rFonts"))
    assert rfonts is not None
    # lock 預設：latin=Times New Roman, cjk=標楷體
    assert rfonts.get(qn("w:ascii")) == "Times New Roman"
    assert rfonts.get(qn("w:hAnsi")) == "Times New Roman"
    assert rfonts.get(qn("w:eastAsia")) == "標楷體"
    assert rfonts.get(qn("w:cs")) == "Times New Roman"


def test_lock_fonts_applied_to_heading1(tmp_path):
    """lock 設定的 h1 font_size 與 bold 應寫入 Heading 1 style。"""
    from docx import Document
    from docx.oxml.ns import qn

    html = "<html><body><h1>T</h1></body></html>"
    out = tmp_path / "h1.docx"
    html_to_docx_direct(html, out, lock_file=DEFAULT_LOCK)
    doc = Document(str(out))
    h1 = doc.styles["Heading 1"]
    rpr = h1.element.find(qn("w:rPr"))
    assert rpr is not None
    rfonts = rpr.find(qn("w:rFonts"))
    assert rfonts is not None
    assert rfonts.get(qn("w:eastAsia")) == "標楷體"
    # size: lock h1.font_size = 18
    sz = rpr.find(qn("w:sz"))
    assert sz is not None
    assert int(sz.get(qn("w:val"))) == 36  # 18pt × 2 (half-points)
    # bold
    assert rpr.find(qn("w:b")) is not None


def test_lock_page_size_and_margins(tmp_path):
    """lock 設定的 page_size=A4 與 margins 應寫入 section。"""
    from docx import Document
    from docx.shared import Cm

    html = "<html><body><p>x</p></body></html>"
    out = tmp_path / "page.docx"
    html_to_docx_direct(html, out, lock_file=DEFAULT_LOCK)
    doc = Document(str(out))
    sec = doc.sections[0]
    # A4 width = 21cm (allow tiny float drift from EMU)
    assert abs(sec.page_width.cm - 21.0) < 0.05
    assert abs(sec.page_height.cm - 29.7) < 0.05
    # Margins from lock: top=2.5cm, bottom=2.5cm, left=3cm, right=2cm
    assert abs(sec.top_margin.cm - 2.5) < 0.05
    assert abs(sec.bottom_margin.cm - 2.5) < 0.05
    assert abs(sec.left_margin.cm - 3.0) < 0.05
    assert abs(sec.right_margin.cm - 2.0) < 0.05


def test_lock_fallback_when_lock_missing(tmp_path):
    """lock 檔不存在時，使用 fallback 設定（不 crash）。"""
    html = "<html><body><h1>Fallback</h1><p>x</p></body></html>"
    out = tmp_path / "fb.docx"
    html_to_docx_direct(
        html, out,
        lock_file=tmp_path / "nope_lock.md",  # 不存在
    )
    assert out.exists()
    assert out.stat().st_size > 1024


# ────────────────────────────────────────────────────────────────────
# 8. Inline element conversion (strong / em / code / a)
# ────────────────────────────────────────────────────────────────────

def test_inline_elements_preserved(tmp_path):
    """<strong>/<em>/<code>/<a> 在 docx 中以對應 runs 保留。"""
    from docx import Document

    html = """<html><body>
<p>這是 <strong>粗體</strong> 與 <em>斜體</em> 還有 <code>monospace</code>。</p>
<p>連結：<a href="https://example.com">Example</a></p>
</body></html>"""
    out = tmp_path / "inline.docx"
    html_to_docx_direct(html, out, lock_file=DEFAULT_LOCK)
    doc = Document(str(out))
    # 收集所有 run 文字
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "粗體" in all_text
    assert "斜體" in all_text
    assert "monospace" in all_text
    assert "Example" in all_text
    # bold + italic 應該有對應的 run
    found_bold = False
    found_italic = False
    for p in doc.paragraphs:
        for r in p.runs:
            if r.bold:
                found_bold = True
            if r.italic:
                found_italic = True
    assert found_bold
    assert found_italic


# ────────────────────────────────────────────────────────────────────
# 9. List conversion
# ────────────────────────────────────────────────────────────────────

def test_ul_ol_converted_to_paragraphs(tmp_path):
    """<ul>/<ol> 會被轉成多個段落（帶項目符號前綴）。"""
    from docx import Document

    html = """<html><body>
<ul><li>第一項</li><li>第二項</li></ul>
<ol><li>步驟一</li><li>步驟二</li></ol>
</body></html>"""
    out = tmp_path / "list.docx"
    html_to_docx_direct(html, out, lock_file=DEFAULT_LOCK)
    doc = Document(str(out))
    text = " ".join(p.text for p in doc.paragraphs)
    assert "第一項" in text
    assert "第二項" in text
    assert "步驟一" in text
    assert "步驟二" in text


# ────────────────────────────────────────────────────────────────────
# 10. Unit tests for helpers
# ────────────────────────────────────────────────────────────────────

def test_to_cm_parses_units():
    assert abs(_to_cm("2.5cm") - 2.5) < 1e-6
    assert abs(_to_cm("25mm") - 2.5) < 1e-6
    assert abs(_to_cm("1in") - 2.54) < 1e-6
    assert abs(_to_cm("100pt") - 3.52778) < 0.01
    assert _to_cm(None) == 2.5  # default
    assert _to_cm(3.0) == 3.0
    assert _to_cm("invalid") == 2.5  # default on bad input


def test_load_lock_with_default_path():
    """_load_lock 預設讀 examples/lock.md，應有 fonts.cjk/latin。"""
    lock = _load_lock(DEFAULT_LOCK)
    assert "fonts" in lock
    assert "cjk" in lock["fonts"]
    assert "latin" in lock["fonts"]


def test_load_lock_fallback_on_missing(tmp_path):
    lock = _load_lock(tmp_path / "nope.md")
    # 預設值
    assert lock["fonts"]["cjk"] == "標楷體"
    assert lock["fonts"]["latin"] == "Times New Roman"


# ────────────────────────────────────────────────────────────────────
# 11. Error handling
# ────────────────────────────────────────────────────────────────────

def test_missing_html_file_raises(tmp_path):
    out = tmp_path / "out.docx"
    with pytest.raises(HTMLToDOCXDirectError):
        html_to_docx_direct(str(tmp_path / "nope.html"), out,
                            lock_file=DEFAULT_LOCK)


def test_empty_html_string_raises(tmp_path):
    out = tmp_path / "out.docx"
    with pytest.raises(HTMLToDOCXDirectError):
        html_to_docx_direct("", out, lock_file=DEFAULT_LOCK)


def test_template_not_found_raises(tmp_path):
    out = tmp_path / "out.docx"
    with pytest.raises(HTMLToDOCXDirectError):
        html_to_docx_direct(
            "<html><body><p>x</p></body></html>",
            out,
            template=tmp_path / "nope_template.docx",
            lock_file=DEFAULT_LOCK,
        )
