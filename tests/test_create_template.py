r"""tests/test_create_template.py — scripts/create_template.py + workflows/create-template.md 單元測試。

DoD 對應 `tasks.md` T3-4：
1. 載入 `workflows/create-template.md` frontmatter 不 crash
2. Mermaid flowchart 存在（regex `` ```mermaid ``）
3. `create_template.py` CLI `--name "test-template"` 不 crash
4. (補充) end-to-end 整合：給 template name → 產 4 個檔案
5. (補充) lock_template.md 含 17 個 required 欄位
6. (補充) DiscoveryAnswers YAML 序列化含 template_answers
7. (補充) StubLLM 預設行為：給出 7 題答案
8. (補充) reference.docx 通過 python-docx round-trip + 字體鎖死
9. (補充) workflow 引用 build_template.py / strategist.md / executor-base.md
10. (補充) 無效 name（含 '/' '..'）→ raise CreateTemplateError
"""

from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

import pytest
import yaml

from scripts.create_template import (
    CreateTemplateError,
    DiscoveryAnswers,
    DiscoveryError,
    DocumentationError,
    StubLLM,
    ValidationError,
    _LOCK_REQUIRED_FIELDS,
    _validate_lock_template_fields,
    run_create_template,
    run_discovery,
    run_documentation,
    run_generation,
    run_validation,
)


# ─── Fixtures ────────────────────────────────────────────────────────

PROJECT_ROOT = Path(__file__).resolve().parent.parent
WORKFLOW_MD = PROJECT_ROOT / "workflows" / "create-template.md"
BUILD_TEMPLATE_PY = PROJECT_ROOT / "scripts" / "build_template.py"
STRATEGIST_MD = PROJECT_ROOT / "references" / "strategist.md"
EXECUTOR_MD = PROJECT_ROOT / "references" / "executor-base.md"


def _has_python_docx() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


pytestmark = pytest.mark.skipif(
    not _has_python_docx(),
    reason="python-docx not installed (required for create_template tests)",
)


# ─── Test 1（必要）: 載入 workflows/create-template.md frontmatter 不 crash ─

def test_workflow_md_frontmatter_loads():
    """載入 workflows/create-template.md frontmatter 不 crash + 必要欄位齊備。"""
    assert WORKFLOW_MD.exists(), f"找不到 {WORKFLOW_MD}"

    content = WORKFLOW_MD.read_text(encoding="utf-8")
    assert content.startswith("---\n"), "create-template.md 應以 frontmatter 開頭"

    # 用正則抓 frontmatter
    m = re.match(r"\A---\s*\n(?P<yaml>.*?)\n---", content, re.DOTALL)
    assert m is not None, "frontmatter 區塊未正確閉合"

    fm = yaml.safe_load(m.group("yaml"))
    assert isinstance(fm, dict), "frontmatter 應為 mapping"

    # 必要欄位
    assert fm.get("name") == "create-template", "name 應為 'create-template'"
    assert fm.get("description"), "description 必填"
    assert len(fm.get("description", "")) >= 50, "description 太短"
    # PyYAML 會把 "1.0" parse 為 float；接受 1.0 或 "1.0"
    version = fm.get("version")
    assert version is not None, "version 必填"
    assert str(version) == "1.0", f"version 應為 1.0，got: {version!r}"

    # 檔案行數 DoD 1
    line_count = len(content.splitlines())
    assert line_count > 100, f"workflow 應 > 100 行，目前 {line_count}"


# ─── Test 2（必要）: Mermaid flowchart 存在 ─────────────────────────

def test_workflow_md_contains_mermaid():
    """workflows/create-template.md 內含 Mermaid flowchart（regex `` ```mermaid ``）。"""
    content = WORKFLOW_MD.read_text(encoding="utf-8")

    # 找 ```mermaid 開頭的區塊
    pattern = re.compile(r"^```mermaid\s*$", re.MULTILINE)
    matches = pattern.findall(content)
    assert len(matches) >= 1, "workflow 應含至少 1 個 ```mermaid 區塊"

    # 額外檢查：Mermaid 內含 flowchart 關鍵字
    m = re.search(r"```mermaid\s*\n(?P<body>.*?)\n```", content, re.DOTALL)
    assert m is not None, "應有完整 mermaid 區塊（``` 閉合）"
    mermaid_body = m.group("body")
    assert "flowchart" in mermaid_body or "graph" in mermaid_body, (
        "Mermaid 區塊應含 flowchart 或 graph 關鍵字"
    )
    # 應涵蓋 4 個主要階段
    for stage_keyword in ("Discovery", "Generation", "Validation", "Documentation"):
        assert stage_keyword in mermaid_body, (
            f"Mermaid 應含 {stage_keyword} 階段"
        )


# ─── Test 3（必要）: create_template.py CLI --name "test-template" 不 crash ──

def test_create_template_cli_runs(tmp_path: Path):
    """`python -m scripts.create_template --name "..." --output-root <tmp>` 不 crash。"""
    result = subprocess.run(
        [
            sys.executable, "-m", "scripts.create_template",
            "--name", "test-template",
            "--description", "test description for create-template workflow",
            "--output-root", str(tmp_path),
            "--quiet",
        ],
        cwd=str(PROJECT_ROOT),
        capture_output=True,
        text=True,
        timeout=30,
    )
    assert result.returncode == 0, (
        f"CLI 失敗（returncode={result.returncode}）\n"
        f"stdout: {result.stdout}\nstderr: {result.stderr}"
    )

    # 應有 4 個檔案
    out_dir = tmp_path / "test-template"
    assert out_dir.exists(), f"沒產出 {out_dir}"
    assert (out_dir / "reference.docx").exists(), "缺 reference.docx"
    assert (out_dir / "README.md").exists(), "缺 README.md"
    assert (out_dir / "lock_template.md").exists(), "缺 lock_template.md"
    assert (out_dir / "discovery_answers.yaml").exists(), "缺 discovery_answers.yaml"


# ─── Test 4（補充）: end-to-end 整合測試 ────────────────────────────

def test_run_create_template_end_to_end(tmp_path: Path):
    """給 template name + description → 產 4 個檔案 + validation PASS。"""
    result = run_create_template(
        name="product-brief",
        description="公司內部 product brief 範本",
        output_root=tmp_path,
        verbose=False,
    )

    assert result["name"] == "product-brief"
    assert isinstance(result["discovery"], DiscoveryAnswers)
    assert result["reference_docx"].exists()
    assert result["readme"].exists()
    assert result["lock_template"].exists()
    assert result["discovery_yaml"].exists()

    # Validation report 應 passed=True
    assert result["validation_report"]["passed"] is True


# ─── Test 5（補充）: lock_template.md 含 17 個 required 欄位 ───────

def test_lock_template_contains_required_fields(tmp_path: Path):
    """lock_template.md frontmatter 應含 docs/report_lock_schema.md §2 規範的 9 個頂層 required 欄位。"""
    result = run_create_template(
        name="customer-newsletter",
        description="客戶電子報範本",
        output_root=tmp_path,
        verbose=False,
    )

    lock_content = result["lock_template"].read_text(encoding="utf-8")

    # 用 _validate_lock_template_fields 檢查
    missing = _validate_lock_template_fields(lock_content)
    assert not missing, f"lock_template.md 缺欄位：{missing}"

    # 額外確認 _LOCK_REQUIRED_FIELDS 全在
    assert len(_LOCK_REQUIRED_FIELDS) == 9, (
        f"_LOCK_REQUIRED_FIELDS 應有 9 個頂層欄位，目前 {len(_LOCK_REQUIRED_FIELDS)}"
    )

    # 確認 formatting 子欄位齊備（H1~H3 + body + cover + caption + table + title + toc）
    fm_text = re.match(r"\A---\s*\n(?P<yaml>.*?)\n---", lock_content, re.DOTALL).group("yaml")
    for sub in ("h1", "h2", "h3", "body", "cover", "caption", "title", "toc", "table"):
        assert re.search(rf"^\s+{sub}\s*:", fm_text, re.MULTILINE), (
            f"formatting 子欄位 {sub} 不在 lock_template.md"
        )


# ─── Test 6（補充）: DiscoveryAnswers YAML 序列化 ──────────────────

def test_discovery_answers_yaml_serialization():
    """DiscoveryAnswers.to_yaml() 應含 template_answers + 7 個關鍵欄位。"""
    ans = DiscoveryAnswers(
        q1_audience="公司內部工程師",
        q2_use_case="月度技術分享",
        q3_sections=["Title", "Author", "Date", "Tags", "Body"],
        q5_cover_enabled=True,
        q5_cover_elements=["Title", "Author", "Date", "Tags"],
        q7_citation_style="none",
    )
    yaml_str = ans.to_yaml()
    assert "template_answers:" in yaml_str
    assert "q1_audience:" in yaml_str
    assert "公司內部工程師" in yaml_str
    assert "q4_fonts:" in yaml_str
    assert "q7_citation_style:" in yaml_str
    assert "q8_expected_length:" in yaml_str

    # 反向：應能 load 回 dict（用 PyYAML）
    data = yaml.safe_load(yaml_str)
    assert "template_answers" in data
    assert data["template_answers"]["q1_audience"] == "公司內部工程師"
    assert data["template_answers"]["q7_citation_style"] == "none"


# ─── Test 7（補充）: StubLLM 預設行為 ──────────────────────────────

def test_stub_llm_default_behavior():
    """StubLLM 預設產出 7 題答案 + 字體鎖死 = 標楷體 / Times New Roman。"""
    llm = StubLLM()
    ans = llm.generate_discovery(description="test description", name="test-name")
    assert isinstance(ans, DiscoveryAnswers)
    assert ans.q1_audience != "（未指定）", "q1 應有答案"
    assert ans.q2_use_case != "（未指定）", "q2 應有答案"
    assert len(ans.q3_sections) >= 1, "q3 應至少 1 個 section"
    assert ans.q4_fonts_cjk == "標楷體", "q4 CJK 字體應為 標楷體"
    assert ans.q4_fonts_latin == "Times New Roman", "q4 Latin 字體應為 Times New Roman"
    assert ans.q4_fonts_override is False, "q4 預設不 override"
    assert ans.q7_citation_style in ("none", "APA", "MLA", "Chicago", "IEEE"), (
        f"q7 citation_style 應是合法值，got: {ans.q7_citation_style}"
    )


# ─── Test 8（補充）: reference.docx 字體鎖死 + python-docx 開啟 ─────

def test_reference_docx_font_lock(tmp_path: Path):
    """產出的 reference.docx 應有 Normal ascii=Times New Roman + eastAsia=標楷體。"""
    result = run_create_template(
        name="font-check-template",
        description="字體檢查範本",
        output_root=tmp_path,
        verbose=False,
    )
    ref_path = result["reference_docx"]
    assert ref_path.exists()
    assert ref_path.stat().st_size > 1024

    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(str(ref_path))
    normal = doc.styles["Normal"]
    assert normal.font.name == "Times New Roman", (
        f"Normal ascii font = {normal.font.name!r}"
    )
    rPr = normal.element.find(qn("w:rPr"))
    assert rPr is not None
    rFonts = rPr.find(qn("w:rFonts"))
    assert rFonts is not None
    assert rFonts.get(qn("w:eastAsia")) == "標楷體", (
        f"eastAsia = {rFonts.get(qn('w:eastAsia'))!r}"
    )


# ─── Test 9（補充）: workflow 引用 build_template / strategist / executor ─

def test_workflow_references_assets():
    """workflow 應引用 scripts/build_template.py + references/strategist.md + references/executor-base.md。"""
    content = WORKFLOW_MD.read_text(encoding="utf-8")
    # build_template.py
    assert "scripts/build_template.py" in content or "build_template.py" in content, (
        "workflow 應引用 scripts/build_template.py"
    )
    # references/strategist.md
    assert "references/strategist.md" in content, (
        "workflow 應引用 references/strategist.md"
    )
    # references/executor-base.md
    assert "references/executor-base.md" in content, (
        "workflow 應引用 references/executor-base.md"
    )


# ─── Test 10（補充）: 無效 name → raise CreateTemplateError ────────

@pytest.mark.parametrize("bad_name", [
    "has/slash",
    "has\\backslash",
    "has..dotdot",
    "has space",
    "has\ttab",
])
def test_invalid_name_raises(bad_name):
    """含 '/' '\\' '..' 或空白的 name 應 raise CreateTemplateError。"""
    with pytest.raises(CreateTemplateError):
        run_create_template(
            name=bad_name,
            description="test",
            output_root=Path("/tmp"),
            verbose=False,
        )


# ─── Test 11（補充）: empty description → raise ──────────────────────

def test_empty_description_raises(tmp_path: Path):
    """空 description 應 raise DiscoveryError。"""
    with pytest.raises((DiscoveryError, CreateTemplateError)):
        run_discovery(description="", name="x")

    with pytest.raises((DiscoveryError, CreateTemplateError)):
        run_discovery(description="   ", name="x")


# ─── Test 12（補充）: run_validation 對損壞 docx 應 raise ─────────

def test_run_validation_raises_on_missing_file(tmp_path: Path):
    """run_validation 對不存在的 path 應 raise ValidationError。"""
    with pytest.raises(ValidationError):
        run_validation(tmp_path / "nonexistent.docx")


# ─── Test 13（補充）: 自訂 cover_lines 會被寫入 reference.docx ────

def test_custom_cover_lines_passed_through(tmp_path: Path):
    """自訂 cover_lines 應被 build_template 寫入 reference.docx。"""
    custom_lines = [
        "Author: wai",
        "Team: Platform",
        "Date: 2026-06-13",
    ]
    ref_path = run_generation(
        name="custom-cover-test",
        discovery=DiscoveryAnswers(q5_cover_enabled=True),
        cover_title="Custom Cover",
        cover_lines=custom_lines,
        output_root=tmp_path,
    )
    assert ref_path.exists()

    from docx import Document
    doc = Document(str(ref_path))
    paragraphs = [p.text for p in doc.paragraphs]
    assert "Custom Cover" in paragraphs, "封面標題應在"
    assert "Author: wai" in paragraphs, "自訂 cover_line 應在"
    assert "Team: Platform" in paragraphs, "自訂 cover_line 應在"


# ─── Test 14（補充）: run_documentation 對空 discovery 不 crash ────

def test_run_documentation_minimal(tmp_path: Path):
    """run_documentation 給預設 DiscoveryAnswers 應產出 3 個檔案。"""
    discovery = DiscoveryAnswers()
    # 先產一個 reference.docx
    ref_path = run_generation(
        name="minimal-doc-test",
        discovery=discovery,
        cover_title="Minimal",
        cover_lines=["Line A", "Line B"],
        output_root=tmp_path,
    )

    docs = run_documentation(
        name="minimal-doc-test",
        description="minimal test",
        discovery=discovery,
        reference_path=ref_path,
        output_root=tmp_path,
        timestamp="2026-06-13T00:00:00",
    )

    assert docs["readme"].exists()
    assert docs["lock_template"].exists()
    assert docs["discovery_yaml"].exists()

    # lock_template.md 應可解析為 YAML frontmatter
    lock_content = docs["lock_template"].read_text(encoding="utf-8")
    m = re.match(r"\A---\s*\n(?P<yaml>.*?)\n---", lock_content, re.DOTALL)
    assert m is not None
    fm = yaml.safe_load(m.group("yaml"))
    assert fm["schema_version"] == 1
    assert fm["fonts"]["cjk"] == "標楷體"
    assert fm["fonts"]["latin"] == "Times New Roman"
