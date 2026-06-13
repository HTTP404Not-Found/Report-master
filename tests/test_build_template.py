# test_build_template.py — tests for scripts.build_template
# 對應 SPEC.md §6.1 R1.1 + tasks.md T2-2
#
# 涵蓋：
#   - 預設輸出存在且 > 1 KB
#   - python-docx 可讀，Normal style 含 Times New Roman
#   - mammoth round-trip 讀得到 Title 或 Heading 1 段落
#   - 各種 --type 都能 build
#   - 自訂封面 title + cover-line 會被寫入

from __future__ import annotations

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from scripts.build_template import (
    build,
    build as build_template,
    COVER_PLACEHOLDERS,
    DEFAULT_CJK_FONT,
    DEFAULT_LATIN_FONT,
    BuildTemplateError,
)


# ────────────────────────────────────────────────────────────────────
# Skip if python-docx missing
# ────────────────────────────────────────────────────────────────────

def _has_python_docx() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


pytestmark = pytest.mark.skipif(
    not _has_python_docx(),
    reason="python-docx not installed (required for build_template tests)",
)


# ────────────────────────────────────────────────────────────────────
# Happy path — core DoD
# ────────────────────────────────────────────────────────────────────

def test_default_output_exists_and_over_1kb(tmp_path):
    out = tmp_path / "tpl.docx"
    result = build(out)
    assert result == out
    assert out.exists()
    assert out.stat().st_size > 1024, (
        f"DOCX size {out.stat().st_size} bytes not > 1KB"
    )


def test_template_readable_by_python_docx_normal_style_has_times_new_roman(tmp_path):
    """python-docx 可讀，Normal style 含 Times New Roman。"""
    from docx import Document

    out = tmp_path / "tpl.docx"
    build(out)

    doc = Document(str(out))
    normal = doc.styles["Normal"]
    # python-docx 高階 API 反映 ascii/hAnsi（不直接含 eastAsia）
    assert normal.font.name == "Times New Roman", (
        f"Normal ascii font = {normal.font.name!r}, expected 'Times New Roman'"
    )
    # 也從 XML 確認 eastAsia
    from docx.oxml.ns import qn
    rPr = normal.element.find(qn("w:rPr"))
    assert rPr is not None, "Normal style 沒有 rPr"
    rFonts = rPr.find(qn("w:rFonts"))
    assert rFonts is not None, "Normal style 沒有 rFonts"
    eastAsia = rFonts.get(qn("w:eastAsia"))
    assert eastAsia == DEFAULT_CJK_FONT, (
        f"Normal eastAsia = {eastAsia!r}, expected {DEFAULT_CJK_FONT!r}"
    )


def test_template_mammoth_roundtrip_finds_title_or_heading(tmp_path):
    """mammoth round-trip 可讀到 Title 或 Heading 1 段落。"""
    import mammoth

    out = tmp_path / "tpl.docx"
    build(out)

    with open(str(out), "rb") as f:
        result = mammoth.extract_raw_text(f)
    text = result.value

    # 預設 academic 封面第一行 = "學術論文範本"
    # Heading 1 = "第一章 緒論"
    has_title = "學術論文範本" in text or "Report-master" in text
    has_heading = "第一章 緒論" in text or "緒論" in text
    assert has_title or has_heading, (
        f"mammoth round-trip 沒看到 Title 或 Heading 1：\n{text}"
    )


# ────────────────────────────────────────────────────────────────────
# Type variants
# ────────────────────────────────────────────────────────────────────

@pytest.mark.parametrize("tpl_type", sorted(COVER_PLACEHOLDERS.keys()))
def test_build_all_types(tmp_path, tpl_type):
    """每種 type 都能 build 且 > 1KB。"""
    out = tmp_path / f"tpl-{tpl_type}.docx"
    build(out, type=tpl_type)
    assert out.exists()
    assert out.stat().st_size > 1024


# ────────────────────────────────────────────────────────────────────
# Cover customization
# ────────────────────────────────────────────────────────────────────

def test_cover_title_override(tmp_path):
    out = tmp_path / "tpl.docx"
    build(out, cover_title="我的專屬封面")
    from docx import Document
    doc = Document(str(out))
    assert doc.paragraphs[0].text == "我的專屬封面"
    assert doc.paragraphs[0].style.name == "Title"


def test_cover_lines_override(tmp_path):
    out = tmp_path / "tpl.docx"
    build(
        out,
        cover_title="X",
        cover_lines=["A: 1", "B: 2", "C: 3"],
    )
    from docx import Document
    doc = Document(str(out))
    paras = [p.text for p in doc.paragraphs]
    assert "X" in paras
    assert "A: 1" in paras
    assert "B: 2" in paras
    assert "C: 3" in paras


# ────────────────────────────────────────────────────────────────────
# Validation integration — template must pass docx_validator
# ────────────────────────────────────────────────────────────────────

def test_template_passes_docx_validator(tmp_path):
    """產出的 DOCX 必須能被 docx_validator 驗通過（字體鎖死）。"""
    from scripts.docx_validator import validate_docx

    out = tmp_path / "tpl.docx"
    build(out)

    rep = validate_docx(str(out))
    assert rep.passed, f"validator issues: {rep.issues}"
    # 確認 cjk_font + latin_font 都 PASS
    cjk_check = next((c for c in rep.checks if c.get("name") == "cjk_font"), None)
    latin_check = next((c for c in rep.checks if c.get("name") == "latin_font"), None)
    assert cjk_check and cjk_check.get("passed"), f"cjk_font failed: {cjk_check}"
    assert latin_check and latin_check.get("passed"), f"latin_font failed: {latin_check}"


# ────────────────────────────────────────────────────────────────────
# Error paths
# ────────────────────────────────────────────────────────────────────

def test_invalid_type_raises():
    with pytest.raises(BuildTemplateError):
        build("/tmp/should_not_exist.docx", type="not_a_real_type")


def test_creates_parent_dir(tmp_path):
    nested = tmp_path / "deep" / "nested" / "tpl.docx"
    build(nested)
    assert nested.exists()